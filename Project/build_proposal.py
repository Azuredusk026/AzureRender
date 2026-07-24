# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as rt

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def format_run(run, font_name="Times New Roman", size_pt=10.5, color_rgb=(0,0,0), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, rt.HYPERLINK, is_external=True)
    
    hyperlink_xml = f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}"/>'
    hyperlink = parse_xml(hyperlink_xml)
    
    new_run_xml = f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    new_run = parse_xml(new_run_xml)
    
    text_node_xml = f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{text}</w:t>'
    text_node = parse_xml(text_node_xml)
    new_run.append(text_node)
    
    rPr_xml = f'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    rPr = parse_xml(rPr_xml)
    if color:
        c_xml = f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{color}"/>'
        c = parse_xml(c_xml)
        rPr.append(c)
    if underline:
        u_xml = f'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>'
        u = parse_xml(u_xml)
        rPr.append(u)
    
    rFonts_xml = f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
    rFonts = parse_xml(rFonts_xml)
    sz_xml = f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="21"/>'
    sz = parse_xml(sz_xml)
    rPr.append(rFonts)
    rPr.append(sz)
    
    new_run.append(rPr)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# Define optimized texts for each section (with precise word counts and addressing feedback)
intro_text = (
    "In contemporary digital media technology and interactive entertainment, real-time Non-Photorealistic Rendering "
    "(NPR) has transcended experimental aesthetics to become a widely adopted paradigm in modern games, animation "
    "pipelines, and interactive media applications. Distinct from photorealistic shading structures that mathematically "
    "simulate real-world energy-preserving light interactions, NPR intentionally prioritizes abstract, expressive "
    "stylized aesthetics, synthesizing hand-drawn anime forms through step-quantized light mapping and explicit contour "
    "boundaries. However, existing NPR studies focus heavily on visual stylization quality (e.g., color quantization, "
    "artistic brushstrokes), while the memory transaction overhead and driver-level execution pipelines of NPR post-processing "
    "models under explicit modern graphics APIs remain underexplored.\n\n"
    "Currently, the vast majority of digital animations, indie games, and cultural heritage visualizations rely unconditionally "
    "on high-level commercial game engines like Unity or Unreal Engine. While these commercial environments lower entry "
    "barriers via abstract visual scripting nodes (e.g., Shader Graph), they fundamentally hide and wrap the critical driver-level "
    "graphics pipeline infrastructure. Consequently, undergraduate researchers and graphics engineers are strictly isolated from "
    "inspecting underlying hardware interactions, explicit synchronization control, and multi-pass buffer lifecycle states.\n\n"
    "This systemic encapsulation and focus on visual features reveal clear technical limitations:\n"
    "1. Research Gap in Explicit NPR Optimizations: While Vulkan subpasses have traditionally been regarded as the primary "
    "mechanism for local attachment reuse, recent Vulkan developments such as Dynamic Rendering Local Read (DRLR) provide "
    "alternative local-memory access strategies. However, empirical studies comparing these approaches under stylized NPR workloads "
    "remain limited. Therefore, it remains unclear whether traditional subpass architectures continue to offer measurable advantages "
    "in memory transaction efficiency and frame-time stability.\n"
    "2. Pipeline Black-Box Obscurity: Commercial engines manage resource transitions and execution synchronization implicitly. "
    "This prevents developers from manually inserting memory barriers or manipulating synchronization primitives across "
    "disparate rendering passes, causing CPU-GPU execution bubbles and pipeline stalls.\n"
    "3. Heavyweight Redundancy: Heavyweight generic engine viewports accommodate physically-based rendering (PBR) operations "
    "by default, introducing heavy runtime feature sets and dynamic data structures that are entirely non-essential and "
    "wasteful for minimalist stylized cartoon pipelines.\n\n"
    "To systematically bridge this research gap, this project explicitly refrains from building a generic commercial game engine. "
    "Instead, it develops a Vulkan-based rendering prototype for comparative evaluation of Vulkan subpasses and Dynamic Rendering Local Read (DRLR). "
    "This software provides a transparent verification sandbox engineered specifically to analyze, profile, and compare explicit "
    "synchronization barriers, custom G-Buffer layouts, and on-chip local-memory caching mechanisms crucial for low-latency modern cartoon rendering. "
    "NPR was selected as the benchmark workload because its reliance on multiple G-buffer reads and screen-space post-processing "
    "operations makes memory-access efficiency a primary performance factor."
)

aims_text = (
    "Aim:\n"
    "To conduct a comparative evaluation of Vulkan subpasses and Dynamic Rendering Local Read (DRLR) for real-time NPR rendering, "
    "focusing on memory transaction overhead and frame-time stability.\n\n"
    "Objectives:\n"
    "O1: To construct a C++17/Vulkan core framework managing device screening, swapchain loops, and allocation pooling to serve as the runtime sandbox.\n"
    "O2: To implement a GLSL Cel Shading and contour detection pipeline over offscreen G-buffers to act as the benchmark workload.\n"
    "O3: To implement a Vulkan subpass-based rendering pipeline.\n"
    "O4: To implement a Dynamic Rendering Local Read (DRLR) pipeline.\n"
    "O5: To compare memory transaction statistics, GPU execution time, frame-time variance, CPU command recording time, and selected hardware counters across all approaches.\n\n"
    "Research Hypotheses:\n"
    "H1: Both Vulkan Subpasses and DRLR reduce memory transaction overhead compared with traditional multi-pass rendering.\n"
    "H2: The relative performance of Vulkan Subpasses and DRLR differs between the selected desktop and mobile test platforms.\n"
    "H3: Reductions in memory transaction overhead are associated with lower frame-time variance during deferred NPR rendering workloads."
)

bg_study_text = (
    "4.1 Vulkan Explicit APIs\n"
    "The transition from legacy, driver-managed graphics frameworks (such as OpenGL) to modern, explicit low-level specifications "
    "represents a major architectural shift in computer graphics pedagogy and engineering (Unterguggenberger et al., 2023). Under "
    "traditional OpenGL configurations, the monolithic graphics driver operates as an implicit state-tracking engine, handling resource "
    "hazard checking, memory swapping, and context management internally (Sellers & Kessenich, 2016). While this hides hardware complexities, "
    "it infers severe CPU-side driver abstraction overhead, resulting in runtime latency spikes and unpredictable optimization hitches (Marques & Ferreira, 2019). "
    "The driver must dynamically validate state changes and resource allocations at runtime, leading to thread serialization and significant CPU driver overhead.\n\n"
    "Vulkan resolves these inefficiencies by shifting memory allocation, pipeline state compilation, and command execution scheduling "
    "directly to the application layer (Sellers & Kessenich, 2016). Under the Vulkan paradigm, commands are explicitly recorded into independent "
    "application-allocated command buffers (Scherl et al., 2020). This allows lock-free, concurrent command recording across multiple CPU worker "
    "threads, eliminating single-threaded CPU driver submission bottlenecks and maximizing CPU core utilization (Scherl et al., 2020). Empirical studies by Marques and Ferreira (2019) "
    "confirm that Vulkan delivers vastly superior frame-time predictability and hardware energy efficiency compared to OpenGL when processing "
    "high-throughput visualization streams. Modern engines leverage this multithreading capability to record draw calls in parallel, eliminating driver bottlenecks.\n\n"
    "4.2 Tile-Based Memory Architectures\n"
    "To optimize multi-pass algorithms, modern GPU hardware leverages Tile-Based Rendering (TBR) or Tile-Based Deferred Rendering (TBDR) "
    "architectures (Kanter, 2016). Instead of rendering the entire frame buffer in a single sweep, TBR divides the screen into small tiles (e.g., 16x16 pixels) "
    "and processes them in high-speed, local on-chip GPU tile memory (Akenine-Möller et al., 2018). In traditional deferred rendering pipelines, intermediate "
    "textures (e.g., G-Buffer data) must be flushed to external Video RAM (VRAM) after creation, only to be read back in a subsequent rendering pass. This round-trip "
    "data movement constitutes a severe memory bandwidth bottleneck. In contrast, Vulkan's subpass architecture allows the GPU to route color, depth, and normal fragments "
    "across successive subpasses directly via transient on-chip tile storage, completely bypassing expensive main VRAM read/write cycles (Danliden & Cederrand, 2020). "
    "By marking attachments as transient using VK_IMAGE_USAGE_TRANSIENT_ATTACHMENT_BIT, memory allocations remain localized to high-speed on-chip cache, "
    "optimizing performance on memory-bandwidth-constrained GPU architectures (Arm, 2021; Kanter, 2016). Subpass dependencies (VkSubpassDependency) explicitly manage "
    "execution and memory synchronization barriers, preventing hardware stalls.\n\n"
    "4.3 Dynamic Rendering and DRLR\n"
    "While subpasses optimize memory bandwidth on tile-based hardware, they introduce significant CPU-side complexity. Subpasses require developers to explicitly declare "
    "render pass attachments, subpass descriptions, and subpass dependencies beforehand, which limits pipeline flexibility. To mitigate this, the Khronos Group introduced the "
    "dynamic rendering extension (VK_KHR_dynamic_rendering), which completely removes the need for VkRenderPass and VkFramebuffer structures, simplifying CPU recording "
    "and command submission. However, traditional dynamic rendering lacked a mechanism for local attachment reuse, forcing all intermediate G-Buffer targets back to VRAM. "
    "To resolve this, the VK_KHR_dynamic_rendering_local_read (DRLR) extension was introduced (Khronos Group, 2023). DRLR allows applications to perform local read operations "
    "from transient attachments within a dynamic rendering setup without requiring explicit render pass structures, while maintaining on-chip tile memory performance. This introduces "
    "a trade-off: subpasses offer rigid, driver-validated synchronization, while DRLR provides simplified CPU overhead and flexible execution at the cost of manual "
    "synchronization management (Khronos Group, 2023). Although both subpasses and DRLR aim to reduce external memory traffic, they differ fundamentally in API structure, "
    "synchronization requirements, and implementation flexibility. Existing literature provides extensive discussion of subpasses, while empirical evaluation of DRLR remains "
    "limited. Consequently, a direct comparison under a controlled NPR workload represents a meaningful research opportunity.\n\n"
    "4.4 NPR as a Memory-Bound Workload\n"
    "While traditional forward rendering pipelines are predominantly arithmetic-bound due to complex fragment shading and light source loops, the deferred Non-Photorealistic Rendering "
    "(NPR) pipeline adopted in this study is expected to be relatively memory-intensive due to repeated G-buffer accesses (Akenine-Möller et al., 2018). Deferred toon shading structures "
    "require the generation of comprehensive Geometry Buffers (G-Buffers) containing normal, depth, and material attributes, followed by subsequent screen-space post-processing passes "
    "for contour extraction and color quantization (Isenberg et al., 2003). In this configuration, each pixel must write to multiple high-precision render targets, and subsequent "
    "edge-detection kernels (e.g., Sobel filters) must read from adjacent pixel attributes in the depth-normal textures. This sequence of writing and reading multiple large attachments "
    "generates a high volume of memory transactions. On standard architectures, flushing these intermediate buffers to external Video RAM (VRAM) creates a severe memory bandwidth bottleneck. "
    "Consequently, deferred NPR is an ideal benchmark workload to analyze the performance characteristics and resource constraints of Vulkan local-memory optimization strategies.\n\n"
    "4.5 Toon Shading\n"
    "Non-Photorealistic Rendering (NPR) departs from the strict radiometric tracking defined by classical photorealistic rendering, focusing "
    "instead on artistic interpretation, abstraction, and edge legibility (Kumar et al., 2019). Unlike Physically-Based Rendering (PBR) "
    "which is predominantly arithmetically bound due to complex microfacet BRDF computations, NPR pipelines that rely on multiple G-buffer "
    "attribute reads and screen-space post-processing convolution kernels may become bandwidth-sensitive under deferred architectures (Akenine-Möller et al., 2018). "
    "Fragment shading equations measure the dot product of the unit surface normal vector (N) and the unit light source vector (L) "
    "to determine diffuse exposure (Akenine-Möller et al., 2018):\n"
    "    fdiffuse = max(N . L, 0.0)\n"
    "In traditional configurations, this continuous scalar value is evaluated via hardcoded step functions to produce discrete lighting tiers "
    "(Kim et al., 2001). Modern real-time production loops scale this model by mapping the dot product value directly to custom lookup textures, "
    "known as Ramp Textures or color gradients (Barla et al., 2006). As detailed by Barla et al. (2006) in their extended X-toon model, utilizing "
    "multidimensional mapping structures allows artists to dynamically control color gradients based on secondary views or environmental depth, "
    "capturing complex artistic details like backlighting and level-of-detail shading. Decaudin (1996) advanced this approach by demonstrating "
    "that integrating view-dependent geometric attributes into the toon shader preserves structural forms under intense light configurations. "
    "In large-scale industry implementations, illustrative rendering models combine customized half-Lambertian diffuse "
    "terms and anisotropic specular lobes to produce a cohesive style across dynamic lighting scenarios (Mitchell et al., 2007). Todo et al. (2007) "
    "proposed locally controllable stylized shading models that allow local light manipulation to highlight geometric features. "
    "Overall, the historical progression of toon shading showcases a transition from rigid, driver-level geometry constraints toward highly flexible, "
    "artist-driven models that run efficiently on parallel hardware. By mapping complex illumination math to dynamic textures and anisotropic lobes, "
    "modern shaders achieve a balance between interactive frame-time limits and highly detailed, stylized visual aesthetics (Todo et al., 2007). "
    "Comprehensive surveys confirm that managing these expressive behaviors remains an active focus in graphics style customization (Wegen et al., 2024).\n\n"
    "4.6 Edge Detection\n"
    "The generation of clear line work and outlines forms an indispensable structural component of NPR, serving to define geometric separation "
    "and spatial clarity (Kumar et al., 2019). Line rendering workflows are generally categorized into two methodologies: geometric silhouette "
    "manipulation (e.g., backside normal hull inflation) and screen-space post-processing (Isenberg et al., 2003; Raskar & Cohen, 1999). Geometric extrusion options "
    "necessitate processing duplicate, inverted model geometry meshes via customized vertex shaders, which scales poorly under dense polygon "
    "counts and fails to capture internal surface lines reliably (Raskar & Cohen, 1999). Image-precision edge detection bypasses these limitations "
    "by framing contour extraction as a discrete mathematical post-processing filter executed over a coherent offscreen data buffer (Saidani et al., 2024).\n\n"
    "By drawing structural scene data to offscreen intermediate targets—commonly G-Buffers consisting of linear depth attachments and world-space "
    "normal buffers—discontinuities can be analyzed via image convolution operators (Lauritzen, 2010). The rendering pipeline applies finite-difference "
    "gradient operators, such as the Sobel edge filter, across a 3x3 pixel neighborhood (Saidani et al., 2024). The dual horizontal (Gx) and "
    "vertical (Gy) Sobel convolution matrices sample neighboring pixel data to identify sharp changes. The calculated total gradient magnitude, "
    "derived via G = sqrt(Gx^2 + Gy^2), is subsequently evaluated against an artist-defined threshold. If the magnitude violates this threshold, "
    "the corresponding pixel is classified as an edge and blended with the final scene color. Recent implementations by Saidani et al. (2024) confirm "
    "that Sobel spatial filtering provides excellent edge preservation and execution efficiency when mapped onto real-time parallel processors. "
    "Classic evaluations note that while screen-space methodologies offer mathematical stability, their ultimate rendering efficiency is highly "
    "dependent on memory bandwidth optimizations and effective layout setups (Schofield et al., 1994). By routing these depth-normal G-Buffer targets "
    "through transient local-memory buffers, memory bandwidth overhead is optimized (Danliden & Cederrand, 2020)."
)

methodology_text = (
    "5.1 Experimental Research Design\n"
    "This study employs a quantitative, experimental comparative research design to evaluate and compare Vulkan subpasses and Dynamic Rendering Local Read (DRLR). "
    "Toon rendering serves as the benchmark workload to stress the pipeline, while on-chip local-memory routing mechanisms represent the primary research focus. "
    "Toon rendering was selected because it combines multiple G-buffer reads, screen-space edge detection, and post-processing operations while maintaining "
    "manageable implementation complexity. The system architecture is divided into three distinct pipeline configurations:\n"
    "- Control Group (Traditional Multi-Pass): A deferred rendering pipeline where G-Buffer data is flushed to VRAM attachments in external memory, "
    "and subsequently read back as textures.\n"
    "- Experimental Group A (Subpass-Based): A single VkRenderPass where transient depth-normal attachments are routed via on-chip GPU tile memory, bypassing VRAM writes.\n"
    "- Experimental Group B (Dynamic Rendering Local Read): A pipeline utilizing VK_KHR_dynamic_rendering_local_read. G-Buffer attachments "
    "use VK_IMAGE_USAGE_TRANSIENT_ATTACHMENT_BIT, enabling local attachment reads within a dynamic rendering workflow that does not require explicit VkRenderPass objects, while allowing implementations to exploit on-chip tile memory when available.\n\n"
    "Limitations of Scope: The study compares DRLR against traditional multi-pass rendering and subpass-based rendering. Pure dynamic rendering without "
    "local read support is outside the project scope, and the findings do not attempt to generalize results to all GPU architectures (e.g., mobile TBDR "
    "vs. desktop immediate-mode architectures).\n\n"
    "5.2 Experimental Variables\n"
    "- Independent Variables: Pipeline configuration (Control, Experimental A, Experimental B), resolution scale (720p, 1080p, 1440p), "
    "mesh complexity (Low, High), and hardware architecture (Desktop IMR vs. Mobile TBR).\n"
    "- Dependent Variables: Estimated memory traffic (measured as memory transaction volume), GPU execution time (milliseconds), CPU command recording time (microseconds), frame-time variance, and selected hardware counters (e.g., GPU cycles and cache hits).\n"
    "- Control Variables: Fixed hardware platforms (Desktop: laptop with NVIDIA RTX 4060; Mobile: Xiaomi 14 with Snapdragon 8 Gen 3 Adreno 750), consistent camera fly-through paths, and identical directional lighting vectors.\n\n"
    "5.3 Benchmark Scene and Measurement Strategy\n"
    "To map standard rendering workloads against the mesh complexity variables, three standard benchmark models are utilized:\n"
    "- Low Complexity (20k vertices): Stanford Bunny mesh.\n"
    "- High Complexity (500k vertices): Stanford Dragon mesh.\n"
    "Resolution is expected to have a larger impact than geometry complexity because the benchmark workload is dominated by screen-space G-buffer processing.\n\n"
    "Performance data will be collected using RenderDoc v1.x and vendor-specific GPU profiling utilities (such as NVIDIA Nsight Graphics or AMD Radeon Developer Tool Suite for desktop, and Snapdragon Profiler or Android GPU Inspector for mobile) "
    "by capturing the pipeline state and analyzing memory transaction statistics to estimate overall memory traffic (Sellers & Kessenich, 2016). CPU command recording times "
    "will be measured using high-precision CPU chronometers. Frame times will be recorded over 1,000 frames to analyze statistical variance and frame-time stability (Marques & Ferreira, 2019). "
    "Selected vendor-specific hardware counters (e.g., GPU active cycles, shader ALU utilization, and cache hits) will be captured using profiling tools to analyze execution efficiency across all rendering modes.\n\n"
    "5.4 Technical Implementation Environment\n"
    "- Graphics API: Vulkan 1.3 + VK_KHR_dynamic_rendering + VK_KHR_dynamic_rendering_local_read (or a Vulkan 1.4-capable implementation; managing queue synchronization, descriptor indexing, and explicit dependencies) (Sellers & Kessenich, 2016).\n"
    "- Support Frameworks: C++17; GLFW v3.x; GLM v0.9.x; ImGui (live parameters calibration).\n\n"
    "5.5 Ethical Research Commitments\n"
    "All imported 3D models will be sourced from public domain databases holding verified CC0 or open-source licenses, preventing copyright violations. "
    "The complete engine source code registry will be hosted transparently on GitHub, supporting open science, reproducibility, and collaborative verification. "
    "All engineering concepts, mathematical logic, and designs sourced from external literature are fully cited under standard APA 7th formatting, avoiding plagiarism."
)

significance_text = (
    "The project offers key benefits across academic, educational, and regional industrial contexts:\n"
    "- Academic Contribution: The project provides a comparative evaluation of Subpasses and DRLR "
    "under stylized NPR workloads. It establishes a rigorous benchmark comparing VRAM-flushed architectures against on-chip tile memory layouts, "
    "contributing to low-level graphics optimization literature (Danliden & Cederrand, 2020).\n"
    "- Reproducibility & Open Science: The project provides a reproducible benchmark framework for evaluating Vulkan local-memory optimization strategies. "
    "The complete engine source code and testing assets will be hosted on GitHub to support cross-hardware validation.\n"
    "- Pedagogical Low-Level Core: By avoiding commercial software abstraction layers, this prototype provides an open-source reference for "
    "explicit graphics teaching, serving as a reference for students studying memory allocation and GPU synchronization (Unterguggenberger et al., 2023).\n"
    "- Regional Industry Application: Stylized NPR mechanics form a core component of Malaysia's expanding mobile gaming and animation sectors. "
    "Developing open rendering pipelines helps regional studios build custom IP while optimizing performance on bandwidth-constrained mobile hardware."
)

outcomes_text = (
    "Upon successful project conclusion, the research will deliver the following quantified deliverables:\n"
    "- Source Code Registry: A clean C++17/Vulkan rendering framework implementing three experimental rendering pipelines, free of external dependencies or engine wrappers, featuring manual memory allocation and command queue synchronization.\n"
    "- Interactive Render Demonstration: A compiled execution application supporting real-time rendering mode hot-swaps, including Photorealistic Lambertian Shading, Pure Quantized Color, Edge-Detection Contours, and Full Composite Toon Shading.\n"
    "- Performance Execution Profile: Comprehensive profiling metrics detailing frame times, memory footprints, and profiling data validating memory transaction efficiency (Danliden & Cederrand, 2020).\n"
    "- Academic Dissertation: A rigorous final thesis tracking the software architecture, technical benchmarks, and evaluation results.\n\n"
    "Even if empirical testing yields marginal memory traffic savings (e.g., <5% on specific architectures), the research remains a valuable academic contribution "
    "by documenting driver-level behavior, hardware limits, and profiling constraints under explicit APIs (Sellers & Kessenich, 2016).\n\n"
    "Concluding Remarks:\n"
    "This final year project shifts toon shading implementation out of application-layer engine scripting down into low-level graphics pipeline engineering (Unterguggenberger et al., 2023). Constructing a standalone multi-pass engine directly using Vulkan provides valuable insights into explicit memory control and direct GPU scheduling (Sellers & Kessenich, 2016). The finished framework successfully combines stylized artistic rendering with hardware-level optimizations, offering a powerful platform for interactive content creation and future low-level graphics research."
)

# Reference definitions as a list of tuples (reference_text, url)
references_data = [
    ("Akenine-Möller, T., Haines, E., & Hoffman, N. (2018). Real-time rendering (4th ed.). A K Peters/CRC Press.", "https://www.realtimerendering.com/"),
    ("Arm. (2021). Vulkan subpasses: Mobile graphics best practices. Arm Developer.", "https://developer.arm.com/documentation/102557/latest/"),
    ("Barla, P., Thollot, J., & Markosian, L. (2006). X-toon: An extended toon shader. Proceedings of the 4th International Symposium on Non-photorealistic Animation and Rendering (NPAR '06), 127–132.", "https://doi.org/10.1145/1124728.1124749"),
    ("Danliden, A., & Cederrand, S. (2020). Multi sub-pass & multi render-target shading in Vulkan: Performance based comparison in real-time (Bachelor's thesis, Linköping University). DiVA Portal.", "http://liu.diva-portal.org/smash/get/diva2:1440212/FULLTEXT01.pdf"),
    ("Decaudin, P. (1996). Cartoon-looking rendering of 3D scenes. Research Report INRIA, 2919.", "https://inria.hal.science/inria-00073797/"),
    ("Isenberg, T., Freudenberg, B., Halper, N., Hecke, M., & Strothotte, T. (2003). A developer's guide to silhouette detection algorithms for real-time environments. IEEE Computer Graphics and Applications, 23(4), 28–37.", "https://doi.org/10.1109/MCG.2003.1210862"),
    ("Kanter, D. (2016). Tile-based deferred shading in modern GPUs. Microprocessor Report, 30(5), 1-12.", "https://www.realworldtech.com/tile-based-deferred-shading/"),
    ("Khronos Group. (2023). Vulkan extension: VK_KHR_dynamic_rendering_local_read. Khronos Registry.", "https://registry.khronos.org/vulkan/specs/1.3-extensions/man/html/VK_KHR_dynamic_rendering_local_read.html"),
    ("Kim, K., Kim, J., & Kyung, C. (2001). A study on the real-time toon rendering for 3D geometry model. Proceedings of the 9th Pacific Conference on Computer Graphics and Applications, 138–147.", "https://doi.org/10.1109/PCCGA.2001.962866"),
    ("Kumar, M., Poornima, B., Nagendraswamy, H. S., Manjunath, C., & Rangaswamy, B. E. (2019). A comprehensive survey on non-photorealistic rendering and benchmark developments for image abstraction and stylization. Iran Journal of Computer Science, 2(3), 131–160.", "https://doi.org/10.1007/s42044-019-00034-1"),
    ("Lauritzen, A. (2010). Deferred rendering for current and future rendering pipelines. SIGGRAPH 2010 Beyond Programmable Shading Course.", "https://www.realtimerendering.com/downloads/DeferredShading_beyond_advanced.pdf"),
    ("Marques, R., & Ferreira, J. C. (2019). Evaluating the performance and energy efficiency of OpenGL and Vulkan on a graphics rendering server. 2019 IEEE 13th International Symposium on Embedded Multicore/Many-core Systems-on-Chip (MCSoC), 57–64.", "https://doi.org/10.1109/MCSoC.2019.00018"),
    ("Mitchell, J., Francke, M., & Eng, D. (2007). Illustrative rendering in Team Fortress 2. Proceedings of the 5th International Symposium on Non-photorealistic Animation and Rendering (NPAR '07), 71–76.", "https://doi.org/10.1145/1274871.1274883"),
    ("Raskar, R., & Cohen, M. (1999). Image precision silhouette edges. Proceedings of the 1999 Symposium on Interactive 3D Graphics (I3D '99), 135–140.", "https://doi.org/10.1145/300523.300539"),
    ("Saidani, T., Ghodhbani, R., Ben Ammar, M., Kouki, M., Algarni, M. H., Said, Y., Kachoukh, A., Alsuwaylimi, A. A., Maqbool, A., & Abd-Elkawy, E. H. (2024). Design and implementation of a real-time image processing system based on Sobel edge detection using model-based design methods. International Journal of Advanced Computer Science and Applications, 15(3).", "https://doi.org/10.14569/IJACSA.2024.0150328"),
    ("Scherl, C., Hamidouche, W., & Déforges, O. (2020). Multithreaded rendering for cross-platform 3D visualization based on Vulkan API. The International Archives of the Photogrammetry, Remote Sensing and Spatial Information Sciences, XLIV-4/W1-2020, 57–64.", "https://doi.org/10.5194/isprs-archives-XLIV-4-W1-2020-57-2020"),
    ("Sellers, G., & Kessenich, J. (2016). Vulkan programming guide: The official guide to learning Vulkan. Addison-Wesley Professional.", "https://www.oreilly.com/library/view/vulkan-programming-guide/9780134495545/"),
    ("Todo, H., Anjyo, K., Baxter, W., & Igarashi, T. (2007). Locally controllable stylized shading. ACM Transactions on Graphics (SIGGRAPH 2007), 26(3), Article 17.", "https://doi.org/10.1145/1276377.1276395"),
    ("Unterguggenberger, J., Kerbl, B., Steinberger, M., Schmalstieg, D., & Wimmer, M. (2023). Vulkan all the way: Transitioning to a modern low-level graphics API in academia. Computers & Graphics, 111, 155–165.", "https://doi.org/10.1016/j.cag.2023.02.001"),
    ("Wegen, O., Scheibel, W., Trapp, M., Richter, R., & Döllner, J. (2024). A survey on non-photorealistic rendering approaches for point cloud visualization. IEEE Transactions on Visualization and Computer Graphics.", "https://doi.org/10.1109/TVCG.2024.34017149")
]

def format_cell_text_with_bold_headings(cell, content, row_idx):
    paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
    for i, p_text in enumerate(paragraphs):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
            
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        
        # Enforce justified paragraph layout
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Check subheadings rules and format runs
        trimmed = p_text.strip()
            
        # Rule 1: Starts with numeric hierarchy headings (e.g. 4.1, 4.2, 5.1, etc.) -> entire line bold
        starts_with_hierarchy = False
        for prefix in ["4.1 ", "4.2 ", "4.3 ", "4.4 ", "4.5 ", "4.6 ", "5.1 ", "5.2 ", "5.3 ", "5.4 ", "5.5 "]:
            if trimmed.startswith(prefix):
                starts_with_hierarchy = True
                break
                
        if starts_with_hierarchy:
            run = p.add_run(p_text)
            format_run(run, font_name="Times New Roman", size_pt=10.5, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Rule 2: Bullet point bold label parsing (e.g. "- Academic Contribution: ...")
        if trimmed.startswith("- ") and ":" in trimmed:
            parts = trimmed.split(":", 1)
            run_bullet = p.add_run("- ")
            format_run(run_bullet, font_name="Times New Roman", size_pt=10.5)
            
            label_text = parts[0][2:].strip() + ":"
            run_label = p.add_run(label_text)
            format_run(run_label, font_name="Times New Roman", size_pt=10.5, bold=True)
            
            run_desc = p.add_run(parts[1])
            format_run(run_desc, font_name="Times New Roman", size_pt=10.5)
            continue
            
        # Rule 3: Direct labels (e.g. "Aim:", "Objectives:", "Research Hypotheses:", "Concluding Remarks:")
        direct_headers = ["Aim:", "Objectives:", "Research Hypotheses:", "Concluding Remarks:", 
                          "1. Research Gap in Explicit NPR Optimizations:", "2. Pipeline Black-Box Obscurity:", "3. Heavyweight Redundancy:"]
        matched_header = None
        for dh in direct_headers:
            if trimmed.startswith(dh):
                matched_header = dh
                break
                
        if matched_header:
            run_hdr = p.add_run(matched_header)
            format_run(run_hdr, font_name="Times New Roman", size_pt=10.5, bold=True)
            
            rest = p_text[len(matched_header):]
            if rest:
                run_rest = p.add_run(rest)
                format_run(run_rest, font_name="Times New Roman", size_pt=10.5)
            if not rest.strip():
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Standard paragraph text
        run = p.add_run(p_text)
        format_run(run, font_name="Times New Roman", size_pt=10.5)

def populate_proposal_table(table):
    # 1. Update Title in Row 1 Cell 1
    cell_title = table.rows[1].cells[1]
    cell_title.text = ""
    p_title = cell_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("Comparative Evaluation of Vulkan Subpasses and Dynamic Rendering Local Read for Real-Time NPR Rendering")
    format_run(run_title, font_name="Times New Roman", size_pt=10.5, bold=True)

    # 2. Populate other sections with bold subheadings and justified formatting
    row_mapping = {
        2: intro_text,
        3: aims_text,
        4: bg_study_text,
        5: methodology_text,
        6: significance_text,
        7: outcomes_text
    }
    
    for row_idx, content in row_mapping.items():
        cell = table.rows[row_idx].cells[1]
        cell.text = ""
        format_cell_text_with_bold_headings(cell, content, row_idx)
        
    # 3. Populate References row (Row 8) with clickable URLs
    cell_ref = table.rows[8].cells[1]
    cell_ref.text = ""
    for idx, (ref_text, url) in enumerate(references_data):
        p = cell_ref.paragraphs[0] if idx == 0 and cell_ref.paragraphs else cell_ref.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add the citation text
        run_ref = p.add_run(ref_text + " ")
        format_run(run_ref, font_name="Times New Roman", size_pt=10.5)
        
        # Add the clickable URL
        add_hyperlink(p, url, url, color="0000FF", underline=True)

def create_gantt_chart_table(cell):
    headers = ["Phase", "Milestone Task Specification", "M1", "M2", "M3", "M4", "M5", "M6", "M7"]
    rows_data = [
        ("P1", "Theoretical Exploration & Core Setup: Vulkan specifications, environment setup, and GLFW window.", "X", "X", "", "", "", "", ""),
        ("P2", "Infrastructural Core Construction: Vulkan instance, device creation, and G-buffer structure setup.", "", "X", "X", "", "", "", ""),
        ("P3", "Baseline Multi-Pass Pipeline: Implementing traditional deferred shading with VRAM-flushed G-buffers.", "", "", "X", "X", "", "", ""),
        ("P4", "Subpass Pipeline Development: Implementing subpass-based deferred shading and transient on-chip attachment routing.", "", "", "", "X", "X", "", ""),
        ("P5", "DRLR Pipeline Development: Implementing dynamic rendering local read pipeline and local descriptor reads.", "", "", "", "", "X", "X", ""),
        ("P6", "Profiling & Analysis: Telemetry profiling (Nsight, Snapdragon Profiler) for memory, power, and times.", "", "", "", "", "", "X", ""),
        ("P7", "Thesis Writing & Evaluation: Drafting final dissertation and compiling experimental benchmarks.", "", "", "", "", "", "X", "X")
    ]
    
    gantt_tbl = cell.add_table(rows=len(rows_data) + 1, cols=9)
    gantt_tbl.autofit = False
    
    hdr_cells = gantt_tbl.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        set_cell_shading(hdr_cells[idx], "2E4053")
        set_cell_margins(hdr_cells[idx], top=80, bottom=80, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        format_run(run, font_name="Arial", size_pt=9.5, color_rgb=(255,255,255), bold=True)
        
    col_widths = [Inches(0.5), Inches(3.0), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35)]
    for row in gantt_tbl.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    for row_idx, data in enumerate(rows_data):
        row_cells = gantt_tbl.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)
            
            p = row_cells[col_idx].paragraphs[0]
            if col_idx >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if text == "X":
                    row_cells[col_idx].text = ""
                    set_cell_shading(row_cells[col_idx], "5DADE2")
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            if len(p.runs) > 0:
                format_run(p.runs[0], font_name="Arial", size_pt=9.0)
                
            border_kwargs = {
                "top": {"sz": 4, "val": "single", "color": "D0D3D4"},
                "bottom": {"sz": 4, "val": "single", "color": "D0D3D4"},
                "left": {"sz": 4, "val": "single", "color": "D0D3D4"},
                "right": {"sz": 4, "val": "single", "color": "D0D3D4"}
            }
            set_cell_border(row_cells[col_idx], **border_kwargs)
            
        if row_idx % 2 == 1:
            for c_idx in range(2):
                set_cell_shading(row_cells[c_idx], "F2F4F4")

def add_appendix_synthesis_matrix(doc):
    p = doc.paragraphs[-1]
    p.text = ""
    p.insert_paragraph_before().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    
    title_run = p.add_run("Appendix: Literature Review Synthesis Matrix")
    format_run(title_run, font_name="Arial", size_pt=14, color_rgb=(46, 64, 83), bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    headers = [
        "Research Source",
        "Core Stylization / API Focus",
        "Framework & Platform",
        "Identified Structural Research Gaps & Pipeline Limitations",
        "Proposed Project Alignment Strategy"
    ]
    
    matrix_data = [
        (
            "Kim et al. (2001)",
            "Discrete lighting level quantization.",
            "Legacy 3D Pipeline Methods.",
            "Developed before modern explicit graphics APIs and therefore does not consider contemporary synchronization mechanisms.",
            "Translates primitive quantization logic into highly dynamic GLSL programmable fragment operations."
        ),
        (
            "Raskar & Cohen (1999)",
            "Image-precision silhouette edge calculation via G-Buffer analysis.",
            "Early Object/Image Space Algorithms.",
            "Primarily focused on earlier rendering architectures with high geometry overhead under dense polygon counts.",
            "Generates geometric properties onto explicit subpass transient layers to execute efficient line sorting."
        ),
        (
            "Mitchell et al. (2007)",
            "Industrial artistic cartoon shading features.",
            "Custom High-Level Commercial Pipelines.",
            "Implementation details are proprietary and tied to application-layer engine contexts; black-box abstraction hooks.",
            "Implements illustrative shading formulations within a completely open-source, bare-metal Vulkan 1.3 framework."
        ),
        (
            "Danliden & Cederrand (2020)",
            "Low-level Vulkan subpass performance profiling.",
            "Explicit Modern APIs (Vulkan Core Profile).",
            "Focuses entirely on raw performance architectures; does not address artistic, stylized NPR pipelines.",
            "Uses optimized multi-subpass transient memory loops to implement an efficient, low-latency toon rendering pipeline."
        ),
        (
            "Khronos Group (2023)",
            "Dynamic Rendering Local Read (DRLR).",
            "Modern Vulkan Extensions.",
            "Lack of comparative evaluation against traditional subpass-based local-memory mechanisms.",
            "Provides direct experimental comparison between DRLR and Subpass architectures under NPR workloads."
        )
    ]
    
    table = doc.add_table(rows=len(matrix_data) + 1, cols=5)
    table.autofit = False
    
    col_widths = [Inches(1.2), Inches(1.3), Inches(1.3), Inches(2.2), Inches(2.0)]
    
    hdr_cells = table.rows[0].cells
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        set_cell_shading(hdr_cells[idx], "2E4053")
        set_cell_margins(hdr_cells[idx], top=80, bottom=80, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        format_run(run, font_name="Arial", size_pt=9.5, color_rgb=(255,255,255), bold=True)
        
    for row_idx, data in enumerate(matrix_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            
            if len(p.runs) > 0:
                format_run(p.runs[0], font_name="Times New Roman", size_pt=9.5)
                
            border_kwargs = {
                "top": {"sz": 4, "val": "single", "color": "D0D3D4"},
                "bottom": {"sz": 4, "val": "single", "color": "D0D3D4"},
                "left": {"sz": 4, "val": "single", "color": "D0D3D4"},
                "right": {"sz": 4, "val": "single", "color": "D0D3D4"}
            }
            set_cell_border(row_cells[col_idx], **border_kwargs)
            
        if row_idx % 2 == 1:
            for c_idx in range(5):
                set_cell_shading(row_cells[c_idx], "F8F9F9")
                
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

def main():
    doc_path = "FYP1.docx"
    doc = docx.Document(doc_path)
    
    print("Loading", doc_path)
    table = doc.tables[4]
    
    print("Populating standard sections with bold subheadings and hyperlinks...")
    populate_proposal_table(table)
    
    print("Populating Gantt Chart...")
    gantt_cell = table.rows[9].cells[1]
    gantt_cell.text = ""
    p_desc = gantt_cell.paragraphs[0]
    p_desc.paragraph_format.space_after = Pt(6)
    run_desc = p_desc.add_run(
        "The project development timeline spans a 7-month research cycle divided into five primary engineering phases as illustrated in the timeline matrix below:"
    )
    format_run(run_desc, font_name="Times New Roman", size_pt=10.5)
    create_gantt_chart_table(gantt_cell)
    
    print("Adding Literature Review Synthesis Matrix...")
    add_appendix_synthesis_matrix(doc)
    
    out_path = "DMT2309242-WuChenfeng-Proposal_v8.docx"
    doc.save(out_path)
    print("Success! Saved updated document to", out_path)

if __name__ == "__main__":
    main()
