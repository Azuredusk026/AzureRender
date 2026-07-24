# -*- coding: utf-8 -*-
"""Build the detailed FYP development plan as a polished DOCX.

The document is derived from:
  - ../DMT2309242-WuChenfeng-Proposal.docx
  - ../AfterglowRender (local source snapshot)
  - current Khronos Vulkan documentation
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "FYP_Development_Plan_v1.3.docx"
QA_DIR = ROOT / "Project" / ".docx_qa"
ARCHITECTURE_IMAGE = QA_DIR / "fyp_architecture.png"

SKILL_ROOT = Path(
    r"C:\Users\23587\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.723.12215\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, audit_docx_tables  # noqa: E402


# compact_reference_guide preset tokens
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
INK = "18222D"
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF4EA"
PALE_GOLD = "FFF5D9"
PALE_RED = "FCEBEC"
BORDER = "C9D2DC"
WHITE = "FFFFFF"
GREEN = "2E6B3E"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_east_asia_font(run, latin=FONT_LATIN, cjk=FONT_CJK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:cs"), latin)


def set_run(
    run,
    *,
    size=None,
    bold=None,
    italic=None,
    color=INK,
    latin=FONT_LATIN,
    cjk=FONT_CJK,
):
    set_east_asia_font(run, latin, cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color=BORDER, size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(
    cell,
    text,
    *,
    size=9,
    bold=False,
    color=INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line=1.08,
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = line
    r = p.add_run(str(text))
    set_run(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return p


def set_fixed_table_width(table, widths_dxa):
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=sum(widths_dxa),
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS_DXA,
    )


def add_table(
    doc,
    headers,
    rows,
    widths_dxa,
    *,
    font_size=8.8,
    header_fill=LIGHT_BLUE,
    alternating=True,
    alignments=None,
    status_color_column=None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    header_cells = table.rows[0].cells
    for idx, heading in enumerate(headers):
        header_p = set_cell_text(
            header_cells[idx],
            heading,
            size=9,
            bold=True,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.0,
        )
        set_paragraph_keep(header_p, keep_next=True, keep_lines=True)
        shade_cell(header_cells[idx], header_fill)
        set_cell_borders(header_cells[idx])
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            align = (
                alignments[col_idx]
                if alignments and col_idx < len(alignments)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_cell_text(
                cells[col_idx],
                value,
                size=font_size,
                align=align,
                line=1.07,
            )
            fill = WHITE if not alternating or row_idx % 2 == 0 else "F8FAFC"
            if status_color_column is not None and col_idx == status_color_column:
                lower = str(value).lower()
                if any(token in lower for token in ("高", "block", "必须", "p0")):
                    fill = PALE_RED
                elif any(token in lower for token in ("中", "条件", "p1")):
                    fill = PALE_GOLD
                elif any(token in lower for token in ("低", "可选", "p2", "完成")):
                    fill = PALE_GREEN
            shade_cell(cells[col_idx], fill)
            set_cell_borders(cells[col_idx])
        prevent_row_split(table.rows[-1])

    set_fixed_table_width(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_field(run, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    rpr.extend([rfonts, color, underline, size])
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Lead Callout" not in styles:
        style = styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Lead Callout"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor.from_string(NAVY)
    style.font.bold = True
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    style.paragraph_format.left_indent = Inches(0.18)
    style.paragraph_format.right_indent = Inches(0.18)
    style.paragraph_format.space_before = Pt(5)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.2
    ppr = style._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_BLUE)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    ppr.append(borders)

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.font.name = FONT_LATIN
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.font.italic = True
    note._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    note.paragraph_format.space_after = Pt(5)
    note.paragraph_format.line_spacing = 1.1


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    return section


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("DMT2309242 · FYP DEVELOPMENT PLAN · v1.3")
    set_run(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
    )
    r = p.add_run("Wu Chenfeng · 2026–2027")
    set_run(r, size=8.5, color=MUTED)
    p.add_run("\t")
    r = p.add_run("Page ")
    set_run(r, size=8.5, color=MUTED)
    page_run = p.add_run()
    set_run(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(p, keep_next=True, keep_lines=True)
    return p


def add_body(doc, text, *, bold_lead=None, italic=False, color=INK):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, size=11, bold=True, color=color)
        rest = text[len(bold_lead) :]
        if rest:
            r = p.add_run(rest)
            set_run(r, size=11, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run(r, size=11, italic=italic, color=color)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph(style="Lead Callout")
    r = p.add_run(f"{label}　")
    set_run(r, size=10.5, bold=True, color=BLUE)
    r = p.add_run(text)
    set_run(r, size=10.5, bold=True, color=NAVY)
    set_paragraph_keep(p, keep_lines=True)
    return p


def next_numbering_ids(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc, *, kind="bullet", marker="●", start=1):
    abstract_id, num_id = next_numbering_ids(doc)
    numbering = doc.part.numbering_part.element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    lvl.append(start_el)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.append(fonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list(doc, items, *, numbered=False, marker="●", compact=False):
    num_id = create_numbering(
        doc, kind="decimal" if numbered else "bullet", marker=marker
    )
    for item in items:
        p = doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        ppr.append(num_pr)
        p.paragraph_format.left_indent = Twips(540)
        p.paragraph_format.first_line_indent = Twips(-270)
        p.paragraph_format.space_after = Pt(3 if compact else 4)
        p.paragraph_format.line_spacing = 1.17 if compact else 1.25
        if isinstance(item, tuple):
            lead, detail = item
            r = p.add_run(lead)
            set_run(r, size=10.7, bold=True, color=NAVY)
            r = p.add_run(detail)
            set_run(r, size=10.7, color=INK)
        else:
            r = p.add_run(str(item))
            set_run(r, size=10.7, color=INK)
        set_paragraph_keep(p, keep_lines=True)
    return num_id


def draw_architecture_figure(path):
    from PIL import Image, ImageDraw, ImageFont

    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 1040), "white")
    draw = ImageDraw.Draw(image)

    def font(size, bold=False):
        candidates = [
            r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                return ImageFont.truetype(candidate, size)
        return ImageFont.load_default()

    f_title = font(44, True)
    f_layer = font(30, True)
    f_box = font(25, True)
    f_small = font(21, False)
    f_note = font(23, True)

    draw.text((70, 35), "建议的 FYP Benchmark 分层架构", font=f_title, fill=f"#{NAVY}")

    layer_specs = [
        ("应用与实验层", ["CLI / ImGui", "固定相机轨迹", "试验矩阵", "CSV + JSON"], "E8EEF5"),
        ("共享 NPR 工作负载", ["G-buffer", "Toon Lighting", "Contour Pass", "Composite"], "F3F7FB"),
        ("按节点扩展的渲染路径", ["Multi-pass · FIRST", "Subpasses", "DRLR"], "FFF5D9"),
        ("Vulkan Core", ["Device / Features", "Images / Buffers", "Sync2 / Queries", "Pipelines"], "EAF4EA"),
        ("平台适配层", ["Desktop: GLFW", "Android: GameActivity", "文件 / 时间 / 日志"], "F2F4F7"),
    ]

    y = 125
    for label, boxes, fill in layer_specs:
        draw.rounded_rectangle(
            (70, y, 1530, y + 145),
            radius=24,
            fill=f"#{fill}",
            outline=f"#{BORDER}",
            width=3,
        )
        draw.text((95, y + 18), label, font=f_layer, fill=f"#{DARK_BLUE}")
        x = 410
        available = 1090
        gap = 18
        width = int((available - gap * (len(boxes) - 1)) / len(boxes))
        for idx, text in enumerate(boxes):
            x0 = x + idx * (width + gap)
            draw.rounded_rectangle(
                (x0, y + 24, x0 + width, y + 118),
                radius=16,
                fill="white",
                outline=f"#{BLUE if label == '可替换渲染路径' else BORDER}",
                width=3,
            )
            bbox = draw.textbbox((0, 0), text, font=f_box)
            tx = x0 + (width - (bbox[2] - bbox[0])) / 2
            draw.text((tx, y + 42), text, font=f_box, fill=f"#{INK}")
        if y < 705:
            draw.line((800, y + 145, 800, y + 165), fill=f"#{MUTED}", width=4)
            draw.polygon([(790, y + 158), (810, y + 158), (800, y + 172)], fill=f"#{MUTED}")
        y += 170

    draw.rounded_rectangle(
        (70, 955, 1530, 1010),
        radius=16,
        fill=f"#{NAVY}",
    )
    note = "共享不变量：资产、相机、shader 逻辑、格式、分辨率、warm-up、采样帧数与统计流程"
    bbox = draw.textbbox((0, 0), note, font=f_note)
    draw.text(
        ((1600 - (bbox[2] - bbox[0])) / 2, 968),
        note,
        font=f_note,
        fill="white",
    )
    image.save(path, quality=95)


def add_cover(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FINAL YEAR PROJECT · DEVELOPMENT ROADMAP")
    set_run(r, size=10.5, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("FYP 开发实施计划")
    set_run(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "Comparative Evaluation of Vulkan Subpasses and\n"
        "Dynamic Rendering Local Read for Real-Time NPR Rendering"
    )
    set_run(r, size=15, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(56)
    r = p.add_run("基于 proposal 与 AfterglowRender 代码审查制定")
    set_run(r, size=11, italic=True, color=MUTED)

    rows = [
        ("Student", "Wu Chenfeng"),
        ("Student ID", "DMT2309242"),
        ("Plan Version", "v1.3"),
        ("Prepared", "24 July 2026"),
        ("Final Deadline", "25 December 2026"),
        ("Planning Model", "Dependency-driven nodes; no duration estimates"),
        ("Primary Workspace", str(ROOT)),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for row_idx, (label, value) in enumerate(rows):
        set_cell_text(table.rows[row_idx].cells[0], label, size=9.5, bold=True, color=NAVY)
        set_cell_text(table.rows[row_idx].cells[1], value, size=9.5, color=INK)
        shade_cell(table.rows[row_idx].cells[0], LIGHT_BLUE)
        shade_cell(table.rows[row_idx].cells[1], WHITE)
        set_cell_borders(table.rows[row_idx].cells[0])
        set_cell_borders(table.rows[row_idx].cells[1])
    set_fixed_table_width(table, [2160, 7200])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Scope: portfolio-first Vulkan renderer · controlled benchmark · desktop + Android")
    set_run(r, size=9.5, bold=True, color=MUTED)
    add_page_break(doc)


def build_document():
    doc = Document()
    setup_page(doc)
    configure_styles(doc)
    configure_header_footer(doc.sections[0])
    draw_architecture_figure(ARCHITECTURE_IMAGE)
    add_cover(doc)

    add_heading(doc, "文档使用说明", 1)
    add_callout(
        doc,
        "核心建议",
        "不要把 AfterglowRender 整体 fork 成主工程。以 Project/MyVulkanApp 为起点重建一个小型、可测量、跨平台的 benchmark core；"
        "仅借鉴 Afterglow 的 Vulkan 对象封装、subpass/input-attachment 组织、shader 资源声明与调试思路。",
    )
    add_body(
        doc,
        "本计划把 proposal 的研究问题转换为按依赖关系排列的开发节点、决策点、实验门禁和论文证据。"
        "除最终 DDL（2026 年 12 月 25 日）外，不分配周数或工期；完成一个节点的退出条件后再进入下一个节点。",
    )
    add_heading(doc, "计划中的四条原则", 2)
    add_list(
        doc,
        [
            ("作品集优先：", "第一条完整渲染路径固定为传统 Multi-pass；先做成可展示的 Desktop NPR renderer，再扩展 Subpass、DRLR 与 Android。"),
            ("能力真实可证：", "在进入 DRLR 节点前确认 RTX 4060 与 Xiaomi 14 的 Vulkan 版本、DRLR feature、depth/stencil local-read 属性和 profiler 可用性。"),
            ("先做统一输出：", "三个渲染路径必须共享场景、相机、G-buffer 格式和 shader 数学；性能优化不能先于视觉等价。"),
            ("先测核心段，再测整帧：", "把本地 attachment reuse 的 GPU 区段单独计时，同时保留端到端 frame time，避免公共 contour pass 稀释差异。"),
            ("原始数据不可手改：", "程序直接输出逐帧 CSV 与 run metadata JSON；统计脚本只读取原始文件并生成派生表与图。"),
        ],
    )
    add_heading(doc, "文档导航", 2)
    add_table(
        doc,
        ["部分", "用途"],
        [
            ("1–3", "确认研究目标、范围、基线与技术架构"),
            ("4", "按依赖顺序执行开发节点；第一完整路径是 Multi-pass 作品集 renderer"),
            ("5", "处理必须由你确认的产品、作品集与研究决策"),
            ("6–8", "完成实验设计、测试和跨平台实施"),
            ("9–11", "管理风险、开发节奏、论文与演示"),
            ("附录", "复用边界、目录、验收清单与技术来源"),
        ],
        [1800, 7560],
        font_size=9.2,
    )

    add_heading(doc, "1. 项目目标与成功定义", 1)
    add_heading(doc, "1.1 从 proposal 转换出的研究问题", 2)
    add_body(
        doc,
        "研究问题应保持为：在完全相同的 deferred NPR 工作负载下，传统 multi-pass、Vulkan subpasses 与 "
        "Dynamic Rendering Local Read（DRLR）在 attachment 数据流、GPU 时间、CPU command recording 时间和 frame-time "
        "稳定性方面是否存在可重复、平台相关的差异？",
    )
    add_heading(doc, "1.2 交付物", 2)
    deliverable_rows = [
        ("D0", "Portfolio renderer", "Desktop Multi-pass NPR renderer 可独立运行；含完整场景、toon、contour、debug views、README、截图与演示视频。", "P0"),
        ("D1", "Benchmark application", "Windows desktop 与 Android 可运行；至少 desktop 支持三条路径，Android 依据 feature probe 支持可用路径。", "P0"),
        ("D2", "NPR demonstration", "G-buffer、toon lighting、screen-space contour、mode switching、固定相机轨迹。", "P0"),
        ("D3", "Experiment package", "配置、原始 CSV、metadata JSON、统计脚本、图表和可复现实验说明。", "P0"),
        ("D4", "Technical documentation", "架构、同步、attachment layout、feature matrix、依赖许可证与构建说明。", "P0"),
        ("D5", "Thesis + presentation", "方法、结果、限制、结论、poster/slide、录屏与现场演示 fallback。", "P0"),
        ("D6", "Optional polish", "更丰富资产、额外分辨率、GUI 美化、第二款移动设备。", "可选"),
    ]
    add_table(
        doc,
        ["ID", "交付物", "完成条件", "优先级"],
        deliverable_rows,
        [600, 1800, 6060, 900],
        font_size=8.6,
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        status_color_column=3,
    )
    add_heading(doc, "1.3 项目级 Definition of Done", 2)
    add_list(
        doc,
        [
            "三个路径在相同输入下输出视觉等价；自动 screenshot diff 达到预设阈值，且差异原因有记录。",
            "Validation Layer 在所有测试场景无 error；允许的 warning 必须进入 waiver 文件并解释。",
            "每个 benchmark run 都带设备、驱动、Vulkan API、extension/feature、commit、配置、温度/电源条件等 metadata。",
            "每项核心结论可从原始数据重新生成；图表不得只存在于论文文档中。",
            "代码、第三方依赖和资产均有许可证记录；Afterglow 派生代码保留 MIT notice 与来源注释。",
            "最终演示可以离线运行；若移动设备或 profiler 现场失败，可使用固定版本录屏与已验证数据。",
        ],
        marker="□",
    )

    add_heading(doc, "2. 工作区与 AfterglowRender 评估", 1)
    add_heading(doc, "2.1 当前工作区基线", 2)
    baseline_rows = [
        ("Proposal", "DMT2309242-WuChenfeng-Proposal.docx", "研究范围已明确；三路径、双平台、1000-frame 采样和核心指标已定义。"),
        ("Starter app", "Project/MyVulkanApp", "仅有 GLFW window skeleton；Vulkan 初始化尚未实现。"),
        ("Learning source", "Project/Vulkan-Tutorial", "含 desktop、Android、profile、multi-object 等教程，可作为 API 查证与小样本来源。"),
        ("Reference renderer", "AfterglowRender", "MIT；约 438 cpp、438 headers、136 shader files；已有 render pass/subpass/input attachment 抽象。"),
        ("Build state", "Workspace root", "根目录不是 Git repository；应先初始化主工程版本控制，并把大体积参考仓库与数据分区管理。"),
    ]
    add_table(
        doc,
        ["对象", "位置", "判断"],
        baseline_rows,
        [1500, 2760, 5100],
        font_size=8.8,
    )
    add_heading(doc, "2.2 不建议直接基于 AfterglowRender 的原因", 2)
    add_list(
        doc,
        [
            ("范围失配：", "Afterglow 是通用 raster/compute 开发框架，包含 ECS、material、asset hot reload、ACES、Bloom、boids、terrain 等；FYP 只需要一个受控 benchmark。"),
            ("平台失配：", "当前工程是 Visual Studio / Windows 布局，依赖 GLFW、DXC、OpenImageIO、Assimp、ImGui 等桌面库；Android 端迁移成本高。"),
            ("语言与 shader 失配：", "Afterglow 使用 C++20 与 HLSL/DXC，而 proposal 指定 C++17 与 GLSL；直接复用会改变方法章节与工具链。"),
            ("实验可解释性风险：", "大量通用系统会增加 CPU、内存与 command recording 噪声，使三条渲染路径的差异更难归因。"),
            ("DRLR 缺口：", "本地代码已有传统 VkRenderPass/subpass/input-attachment 结构，但未发现 DRLR feature chain 与 dynamic rendering local-read 路径。"),
        ],
    )
    add_heading(doc, "2.3 推荐复用边界", 2)
    reuse_rows = [
        ("AfterglowSubpassContext / RenderPass", "参考设计", "借鉴 attachment description、input/color reference、dependency 构造；用更小的 BenchmarkRenderGraph 重写。"),
        ("Device / PhysicalDevice", "参考并重写", "保留 device screening 思路；必须改成 Features2 pNext chain 并输出完整 capability JSON。"),
        ("Material / Asset / ECS", "不复用", "对 benchmark 过重；改用固定 pipeline + 简单 scene/model loader。"),
        ("Toon shaders", "仅参考视觉", "Afterglow 目录虽有 Toon，但 FYP 应保留统一 GLSL/SPIR-V shader source 和可审计公式。"),
        ("Command / buffer / image wrappers", "选择性借鉴", "只吸收 RAII 生命周期与命名方式，不复制 proxy/reflection 等复杂基础设施。"),
        ("GUI / hot reload", "可选", "作品集节点仅加入必要的 debug UI；hot reload 在 benchmark 稳定后再决定，正式测量时必须关闭或剔除。"),
        ("MIT source", "可合法复用", "复制任何 substantial code 时保留 LICENSE、文件级来源和修改说明。"),
    ]
    add_table(
        doc,
        ["模块", "处理方式", "具体动作"],
        reuse_rows,
        [2700, 1380, 5280],
        font_size=8.6,
    )
    add_callout(
        doc,
        "架构决策 AD-001",
        "主工程采用 clean benchmark implementation；AfterglowRender 保持只读 reference snapshot。只有在某个小模块能显著降低实现复杂度、"
        "且不引入 Windows/C++20/大依赖耦合时才复制，并记录来源。",
    )

    add_heading(doc, "3. 技术范围与架构", 1)
    add_heading(doc, "3.1 Scope boundary", 2)
    scope_rows = [
        ("In scope", "Vulkan instance/device/swapchain；同步；images/buffers；G-buffer；toon lighting；contour；三渲染路径；feature probe；timestamp；CSV/JSON；Windows/Android。"),
        ("Out of scope", "通用 game engine、完整 ECS、动画系统、PBR、shadow、physics、editor、网络、复杂 asset pipeline、新 NPR 风格研究。"),
        ("研究控制", "资产、相机、灯光、shader 数学、attachment 格式、分辨率、帧数、warm-up、编译优化与统计流程保持一致。"),
        ("扩展项", "更多 GPU、MSAA、第二种 contour、VMA/allocator 比较、额外功耗指标，仅在必需节点全部通过后加入。"),
    ]
    add_table(
        doc,
        ["边界", "内容"],
        scope_rows,
        [1740, 7620],
        font_size=9,
    )
    add_heading(doc, "3.2 建议架构", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    picture = run.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.35))
    picture._inline.docPr.set(
        "descr",
        "Layered architecture for the FYP benchmark: application and experiment layer, "
        "shared NPR workload, interchangeable Multi-pass/Subpass/DRLR paths, Vulkan core, "
        "and Desktop/Android platform adapters.",
    )
    picture._inline.docPr.set("title", "Recommended FYP benchmark architecture")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run("Figure 1. 三条渲染路径共享相同 workload 与测量层。")
    set_run(r, size=9, italic=True, color=MUTED)

    add_heading(doc, "3.3 关键接口", 2)
    interface_rows = [
        ("RenderPath", "prepare(), record(), resize(), destroy(), name(), capabilities()", "保证三路径具有同一调用边界。"),
        ("FrameResources", "per-frame command pool/buffer, fences, semaphores, query pool", "避免不同路径使用不同 frames-in-flight 结构。"),
        ("GBufferLayout", "formats, usage flags, clear/load/store policy", "所有路径从一份配置生成 attachments。"),
        ("BenchmarkRunner", "warmup, sample, repeat, randomize, export", "消除手工测试与遗漏 metadata。"),
        ("CapabilityReport", "API/driver/extensions/features/properties/formats", "第 1 个程序输出；在任何 DRLR 开发前运行。"),
        ("SceneScript", "fixed camera path, deterministic time, scene complexity", "避免用户输入与动画时间差引入噪声。"),
    ]
    add_table(
        doc,
        ["接口", "最小职责", "实验价值"],
        interface_rows,
        [1800, 3960, 3600],
        font_size=8.7,
    )

    add_heading(doc, "3.4 建议目录结构", 2)
    directory_rows = [
        ("CMakeLists.txt / CMakePresets.json", "统一 Desktop、Android、Debug、Release、Benchmark presets"),
        ("src/app", "CLI、mode switch、benchmark state machine"),
        ("src/platform/desktop", "GLFW window/surface、文件系统"),
        ("src/platform/android", "GameActivity/NativeActivity、ANativeWindow、assets"),
        ("src/vk", "instance、device、swapchain、resources、sync、query、debug labels"),
        ("src/render/common", "scene、G-buffer、pipelines、shared descriptors"),
        ("src/render/multipass", "传统 multi-pass control"),
        ("src/render/subpass", "VkRenderPass + input attachments"),
        ("src/render/drlr", "dynamic rendering + local read + self-dependency"),
        ("shaders", "GLSL source + generated SPIR-V；三路径共享 include"),
        ("src/benchmark / tools", "runner、CSV/JSON、analysis/plot scripts"),
        ("assets / licenses", "固定测试资产、哈希、来源和许可证"),
        ("tests / docs", "unit/integration/golden tests、ADRs、protocol、results"),
    ]
    add_table(
        doc,
        ["路径", "职责"],
        directory_rows,
        [3240, 6120],
        font_size=8.8,
    )

    add_heading(doc, "3.5 G-buffer 与 pass 拆分", 2)
    gbuffer_rows = [
        ("Normal", "R16G16B16A16_SFLOAT 或经 feature probe 选择的统一格式", "Geometry pass 写；toon 与 contour 读"),
        ("Albedo / material", "R8G8B8A8_UNORM", "Geometry pass 写；toon lighting 读"),
        ("Depth", "D32_SFLOAT（若两平台均支持目标 usage）", "Geometry pass 写；contour 读"),
        ("Toon color", "R8G8B8A8_UNORM", "Lighting pass 写；contour/composite 读"),
        ("Swapchain", "平台支持的 SRGB format", "最终 composite/present"),
    ]
    add_table(
        doc,
        ["Attachment", "建议格式", "生命周期"],
        gbuffer_rows,
        [1740, 3720, 3900],
        font_size=8.8,
    )
    add_callout(
        doc,
        "必须冻结的实验约束",
        "Input attachment / DRLR local read 读取的是同一 framebuffer 坐标。Sobel 3×3 邻域不能伪装成 local-read 工作。"
        "因此核心比较段设为 G-buffer → same-pixel toon lighting；contour 作为三路径共享的独立 sampled-image pass。"
        "分别记录 core-local 段和 full-frame 指标。",
    )
    add_list(
        doc,
        [
            "Multi-pass：结束 geometry render，执行 layout transition / barrier，在独立 lighting render 中按 sampled image 读取 G-buffer。",
            "Subpass：subpass 0 写 G-buffer；framebuffer-space dependency；subpass 1 以 input attachments 做 same-pixel toon lighting。",
            "DRLR：同一次 dynamic rendering 内写 attachments；BY_REGION self-dependency；映射 input attachment indices；切换到 local-read layout；执行 lighting。",
            "Contour：在上述三条路径之后执行完全相同的全屏 pass，使用 normal/depth sampled images 做邻域采样并 composite。",
        ],
    )

    add_heading(doc, "3.6 依赖选择", 2)
    dependency_rows = [
        ("Vulkan SDK / Headers / Loader", "P0", "API 与 validation；锁定版本并记录 sdk/driver。"),
        ("GLFW", "P0 Desktop only", "Windows surface；Android 不使用 GLFW。"),
        ("Android GameActivity + CMake", "P0 Mobile", "ANativeWindow、生命周期和输入。"),
        ("GLM", "P0", "数学；使用一致的 depth/clip-space 宏。"),
        ("glslangValidator 或 shaderc", "P0", "GLSL → SPIR-V；构建时生成，不在 benchmark timing 内编译。"),
        ("tinygltf + stb_image", "P1", "最小资产加载；也可先使用程序生成 mesh。"),
        ("Dear ImGui", "P1 Debug only", "观察配置；Benchmark preset 默认禁用。"),
        ("VMA", "P1", "减少 allocator 工作；如采用则固定 allocation policy，不能在三路径间变化。"),
    ]
    add_table(
        doc,
        ["依赖", "级别", "约束"],
        dependency_rows,
        [2700, 1500, 5160],
        font_size=8.7,
        status_color_column=1,
    )

    add_heading(doc, "4. 无工期开发节点", 1)
    add_callout(
        doc,
        "默认路径决策",
        "第一条完整渲染路径选择传统 Multi-pass。它既是最容易调试和展示的常规 renderer，也是后续 Subpass 与 DRLR 的公平对照基线。"
        "在作品集版本发布前，不并行开发另外两条路径。",
    )
    add_heading(doc, "4.1 节点总览与依赖顺序", 2)
    node_rows = [
        ("N0", "工程与设备基线", "建立可重复 build、Git/许可证结构和 capability_probe；记录 Desktop 与 Xiaomi 14 的真实支持。", "toolchain 固定；capability JSON 保存；关键决策已登记。"),
        ("N1", "Multi-pass 作品集 Renderer", "只做 Desktop；完成类《终末地》的风格化角色与工业科幻场景、G-buffer、toon、contour、composite、debug views。", "角色与场景画面完整；材质分区可信；validation clean。"),
        ("N2", "作品集发布包", "视觉打磨、性能基线、README、架构图、RenderDoc 截图、演示视频和项目描述；受限资产与公开代码分离。", "陌生人可按 README 构建或观看；公开仓库不包含无再分发许可的模型与贴图。"),
        ("N3", "Benchmark-ready 重构", "把 N1 拆成共享 workload、RenderPath 接口、统一资源与计时边界；冻结 golden images。", "Multi-pass 经重构后输出不变；CSV/JSON 首版可用。"),
        ("N4", "Subpass 路径", "实现 VkRenderPass、input attachments、dependencies 与 framebuffer 管理。", "与 Multi-pass 视觉等价；validation clean；独立计时区段存在。"),
        ("N5", "DRLR 路径", "依据 feature probe 实现 dynamic rendering local read；不支持的平台明确标记 NA。", "支持设备可运行；映射、layout 与 self-dependency 有 capture 证据。"),
        ("N6", "Android 路径", "移植共享 core 与可支持路径；处理 GameActivity、surface、asset 和生命周期。", "Xiaomi 14 或替代设备稳定运行；capability truth table 完整。"),
        ("N7", "公平性与自动实验", "统一 shaders/config、执行 cross-path diff，完成 runner、metadata、温控和 profiler SOP。", "pilot matrix 可无人值守执行；异常可恢复；数据 schema 冻结。"),
        ("N8", "正式实验与论文", "执行正式 runs，生成统计与图表，完成 thesis、演示和最终发布包。", "结论可由原始数据重建；提交包在 2026-12-25 前完整。"),
    ]
    add_table(
        doc,
        ["节点", "目标", "主要动作", "退出条件"],
        node_rows,
        [720, 1860, 4080, 2700],
        font_size=8.25,
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
    )

    add_heading(doc, "4.2 N1 · Multi-pass 作品集 Renderer 的实现顺序", 2)
    add_list(
        doc,
        [
            ("建立 Desktop Vulkan core：", "完成 instance、validation、physical/logical device、queues、GLFW surface、swapchain、resize 与 frames-in-flight。"),
            ("完成资源层：", "实现 image/buffer/view/sampler、staging upload、descriptor、pipeline、command 与同步的最小 RAII 封装。"),
            ("确定作品集场景：", "构建冷色工业科幻展示空间，加载角色与环境；首版只要求稳定静态姿态、可控相机、主光和轮廓光。"),
            ("实现 Geometry pass：", "写入 normal、albedo/material 与 depth；提供 G-buffer debug views，先证明数据正确。"),
            ("实现传统 Multi-pass Toon：", "显式 barrier/layout transition 后读取 G-buffer；按 skin、hair、fabric、metal 分配柔化色阶、高光和边缘光。"),
            ("实现 Contour 与 Composite：", "使用低强度 normal/depth screen-space Sobel；避免厚重黑边，完成抗闪烁、曝光、色彩空间和最终合成。"),
            ("加入作品集交互：", "提供相机控制、场景切换、toon/contour 参数、debug view 和性能 overlay；研究 benchmark 模式暂不启用。"),
            ("做视觉与稳定性打磨：", "修复 resize、资源重建、深度精度、边缘噪声、色彩空间和帧间抖动；运行 validation 与 RenderDoc 审查。"),
            ("记录技术证据：", "保存架构图、pass 流程、G-buffer 截图、关键 shader 说明、GPU capture 和已知限制。"),
            ("发布作品集版本：", "生成 Release build、README、功能截图、短视频、技术摘要、依赖许可证与可公开仓库或展示页。"),
        ],
        numbered=True,
    )
    add_callout(
        doc,
        "范围保护",
        "N1/N2 只完成一个高完成度 Desktop Multi-pass renderer。不要在该节点加入 Subpass、DRLR、Android、完整 ECS、复杂 editor、"
        "PBR、动画系统或多种 NPR 风格；这些内容不能阻塞作品集发布。",
    )

    add_heading(doc, "4.3 作品集版本的验收清单", 2)
    add_list(
        doc,
        [
            "一个双击可运行或按 README 可稳定构建的 Windows Release 版本。",
            "一个具有类《终末地》工业科幻语言的完整角色与场景展示，而不是 tutorial triangle 或纯测试房间。",
            "最终画面、Normal、Depth、Albedo、Toon、Contour 等可切换 debug views。",
            "README 说明项目目标、技术栈、渲染流程、你亲自完成的部分、Afterglow 的参考边界及角色/场景资产权利状态。",
            "至少一张 renderer 架构图、一张 pass/G-buffer 图和一组高质量截图。",
            "一段短演示视频：先展示最终画面，再展示 debug views、参数调整和 RenderDoc/性能证据。",
            "Validation Layer 无 error；常规操作、resize 和多次启动不崩溃。",
            "公开仓库不提交《终末地》原始模型、贴图或可恢复资源；提供 placeholder/自制替代资产与本地导入说明。",
            "不提前宣称 Subpass/DRLR 性能结论；作品集版本只展示 Multi-pass renderer 的实现质量。",
        ],
        marker="□",
    )

    add_heading(doc, "4.4 节点执行规则", 2)
    add_list(
        doc,
        [
            "任何时刻只把一个节点设为 Active；新想法进入 backlog，不直接改变当前节点范围。",
            "每个节点拆成可独立验证的 vertical slices；每个 slice 必须留下 build、截图、capture、测试或数据证据。",
            "节点退出条件未满足时，不以“后面再修”为理由进入依赖它的节点。",
            "作品集发布节点 N2 通过后，再把美术展示代码与 benchmark mode 分离，避免展示功能污染正式计时。",
            "最终 DDL 是 2026 年 12 月 25 日；节点顺序固定，但节点内部允许按 vibe 开发方式灵活拆分和迭代。",
        ],
    )

    add_heading(doc, "5. 需要你决定的事项", 1)
    add_body(
        doc,
        "下列事项会明显改变代码结构、作品集呈现或研究范围。U01–U08 已根据当前决定冻结；"
        "U09–U12 保留为后续决策。若实际资产授权状态不同，只修改对应决策记录，不需要重写整份计划。",
    )
    decision_rows = [
        ("U01", "第一渲染路径", "传统 Multi-pass", "N1/N2 只实现一条完整路径；Subpass 与 DRLR 留到作品集发布后。", "已确认"),
        ("U02", "作品集美术方向", "类《终末地》的风格化角色 + 冷色工业科幻场景", "参考图锁定柔和二至三段明暗、材质分区、受控高光与轻轮廓。", "已确认"),
        ("U03", "作品集资产来源", "使用现有角色模型与贴图做本地展示", "必须确认来源与许可；公开仓库排除原始模型/贴图，使用 placeholder 或自制替代。", "有限确认"),
        ("U04", "首发平台", "Windows Desktop only", "Android 留到研究节点；作品集不被移动端适配阻塞。", "采用默认"),
        ("U05", "公开方式", "公开代码、Release、截图和视频", "代码可公开；第三方角色/贴图只在获得再分发许可时随仓库发布。", "已确认"),
        ("U06", "作品集交互范围", "最小 ImGui：相机、材质/灯光、toon/contour、debug views、GPU overlay", "ImGui 只在 Portfolio/Debug preset 启用；Benchmark preset 默认关闭。", "已决定"),
        ("U07", "Afterglow 复用程度", "reference-only；只重写设计模式", "不整体 fork；确需复制 MIT 小模块时保留 LICENSE、来源和修改记录。", "已决定"),
        ("U08", "场景与模型格式", "源资产为 UE 5.7 .uasset；私下从 UE 导出 FBX/纹理，运行时优先 glTF 2.0 + tinygltf", "莱万汀先使用静态姿态；保留 Skeleton，但首个作品集节点不实现动画。", "已确认"),
        ("U09", "NPR 视觉定义", "柔化二至三段漫反射 + 材质高光/边缘光 + 轻量 normal/depth contour", "在首个材质球与角色半身测试后冻结具体 ramp、阈值与轮廓参数。", "部分确认"),
        ("U10", "DRLR 设备策略", "先 probe；Xiaomi 14 不支持时把 cell 记为 NA，并寻找替代 Android", "决定是否能借到兼容设备，以及是否接受 mobile 只比较两条路径。", "N5/N6 前"),
        ("U11", "正式实验规模", "保留 proposal 主矩阵；pilot 后只根据方差裁剪", "决定是否保留两场景、两分辨率、十次 repeat 和两平台。", "N7 结束前"),
        ("U12", "论文与作品集分仓", "同一 core、两个 preset；portfolio assets 与 benchmark data 分目录", "决定使用单仓库还是公开作品集仓库 + 私有研究数据仓库。", "N2/N3 之间"),
    ]
    add_table(
        doc,
        ["ID", "需要决定", "当前决定/建议", "执行边界", "状态"],
        decision_rows,
        [600, 1680, 2760, 3540, 780],
        font_size=7.95,
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    add_heading(doc, "5.1 当前可直接采用的默认决定", 2)
    add_list(
        doc,
        [
            "主工程从 Project/MyVulkanApp 演进；AfterglowRender 保持独立、只读参考。",
            "首个完整 renderer 是 Windows Desktop Multi-pass deferred NPR，不先做 Android。",
            "作品集功能与 benchmark 功能使用同一 core，但通过 Portfolio 与 Benchmark presets 隔离。",
            "现有角色模型/贴图作为本地私有测试资产；公开代码仓库只提交资产接口、映射配置与合法占位内容。",
            "使用最小 ImGui 做视觉调参与 debug views；AfterglowRender 只读参考，不整体 fork。",
            "作品集优先完成视觉、稳定性、文档与演示；研究比较功能在作品集发布后继续。",
            "任何未确认的移动端 DRLR 支持都视为未知，不把 unsupported 设备伪装成 supported。",
        ],
        marker="□",
    )
    add_heading(doc, "5.2 已确认的角色与场景渲染方向", 2)
    visual_rows = [
        ("构图与场景", "角色为视觉主体；使用模块化工业展示台、浅灰/冷白背景和少量青蓝灯带，避免先做大规模开放场景。", "先完成全身英雄镜头、半身材质镜头和一张场景广角。"),
        ("光照", "柔和主光 + 冷色环境光 + 可控 rim light；阴影边界略柔化，保持角色脸部和服装层次。", "ImGui 暴露 key/fill/rim 强度、方向和颜色。"),
        ("皮肤与脸部", "低频平滑法线、柔和二段或三段 diffuse、低强度宽高光；避免金属感和过黑阴影。", "首版不实现 face SDF；只有资产提供对应 mask 时再加入。"),
        ("头发", "支持 alpha cutout/two-sided、独立 hair ramp、方向性或窄带高光；控制耳朵和发丝边缘锯齿。", "先验证 alpha、深度写入、排序与 MSAA/抗锯齿策略。"),
        ("服装与金属", "cloth 使用较宽 rough specular；皮革使用更集中高光；metal/accessory 使用独立 mask、冷色反射和轻 emissive。", "材质参数由 material ID/纹理 mask 驱动，不为单个模型写死 shader。"),
        ("轮廓", "以 normal/depth Sobel 为主，颜色接近深蓝灰且透明度较低；只强调剪影和材质边界。", "Portfolio preset 可调；Benchmark preset 冻结统一参数。"),
        ("后处理", "线性 HDR 合成后 tone mapping、gamma 和轻量 color grade；bloom 仅用于发光部件。", "不让 bloom、TAA 或动态曝光阻塞首个可展示版本。"),
        ("作品集证据", "同时展示 Final、Albedo、Normal、Depth、Material ID、Toon 与 Contour，并给出 RenderDoc pass 视图。", "证明 renderer 的技术工作，而不是只展示第三方角色资产。"),
    ]
    add_table(
        doc,
        ["视觉部分", "目标", "实现约束"],
        visual_rows,
        [1680, 4560, 3120],
        font_size=8.45,
    )
    add_callout(
        doc,
        "资产公开边界",
        "“拥有文件”不等于拥有再分发权。除非你能证明模型与贴图允许公开分发，否则它们只存在于本地 assets_private/；"
        "公开仓库仅提供 importer、材质映射配置、占位资产和获取说明。公开项目名建议使用 Vulkan Stylized Character Renderer，"
        "不要暗示官方关联；发布截图/视频前核对权利方条款，并标注非官方、非商业技术研究。",
    )
    add_heading(doc, "5.3 莱万汀初始测试资产记录", 2)
    add_body(
        doc,
        "只读盘点位置：D:\\Epic\\UE Project\\ZMDRender\\Content\\ZMD\\莱万汀。项目由 Unreal Engine 5.7 关联；"
        "该目录共 37 个 .uasset，约 75.6 MB。Vulkan 工程不直接解析 .uasset，转换应在 Unreal Editor 内完成，"
        "导出结果进入工作区外的私有中间目录，再复制到 assets_private/。",
    )
    asset_rows = [
        ("角色主体", "莱万汀.uasset", "约 17.2 MB；作为 Skeletal Mesh 候选，需在 UE 中核对 LOD、顶点属性、材质槽和单位。"),
        ("骨骼", "莱万汀_Skeleton.uasset", "保留骨骼与 skin weights；首版固定 bind/展示姿态，不实现动画播放器。"),
        ("物理资产", "莱万汀_PhysicsAsset.uasset", "首个 renderer 不使用；不导入碰撞体和 ragdoll 数据。"),
        ("材质", "MI/ 下 12 个实例", "body×2、cloth×5、face、hair、iris、brow、hair shadow；导出后映射为 renderer material classes。"),
        ("贴图", "Tex/ 下 20 张", "含 body、face、iris、hair、cloth 与 weapon；D/N/P/E/M/HN 后缀仅作工作假设，通道含义须从 UE 父材质逐项核对。"),
        ("轮廓与展示", "MI_ZMD_Outline_Inst_LWT_Hair、Base_42", "只作视觉与参数参考；轮廓在 Vulkan 中以 normal/depth pass 重写，不复制 UE 材质图。"),
    ]
    add_table(
        doc,
        ["类别", "已发现资产", "首版处理"],
        asset_rows,
        [1500, 2700, 5160],
        font_size=8.35,
    )
    add_callout(
        doc,
        "莱万汀导出契约",
        "Mesh：FBX 2020.2 或 UE 可用的稳定 FBX，保留 normals、tangents、UV、vertex color、skin weights 与 material slots；"
        "Textures：按原分辨率导出 PNG/TGA，颜色贴图标记 sRGB，normal/mask/packed 参数图保持 linear；"
        "Manifest：记录 UE object path、导出文件、材质槽、贴图通道、色彩空间、alpha 模式和 SHA-256。"
        "任何 packed texture 在通道语义确认前不得接入 shader。",
    )

    add_heading(doc, "6. 实验设计与数据方案", 1)
    add_heading(doc, "6.1 实验条件", 2)
    experiment_rows = [
        ("Rendering path", "3 levels", "Multi-pass / Subpass / DRLR（unsupported 设备记为 NA，不能模拟成 supported）"),
        ("Platform", "2 levels", "RTX 4060 laptop / Xiaomi 14；完整记录 driver、OS、API 与 feature"),
        ("Scene complexity", "2 levels", "低复杂度 procedural/bunny；高复杂度 dragon 或合法替代资产"),
        ("Resolution", "2 common levels", "1280×720 与 1920×1080 offscreen/render area；native resolution 仅作附加结果"),
        ("Repeats", "10 per cell", "建议 10 次独立 run；pilot 后根据方差可调整，但调整需在正式实验前冻结"),
        ("Frames", "300 warm-up + 1000 sample", "proposal 的 1000 frames 保留；warm-up 不进入统计"),
    ]
    add_table(
        doc,
        ["因素", "水平", "定义"],
        experiment_rows,
        [1800, 1740, 5820],
        font_size=8.8,
    )
    add_callout(
        doc,
        "推荐正式矩阵",
        "3 paths × 2 platforms × 2 scenes × 2 resolutions × 10 repeats = 最多 240 runs。"
        "如果某设备不支持 DRLR，该设备的 DRLR cells 标记为 NA，并将主比较拆成 desktop 三路径与 mobile 两路径；不要用不同设备补齐同一 cell。",
    )
    add_heading(doc, "6.2 变量与指标", 2)
    metric_rows = [
        ("GPU core-local time", "µs / frame", "geometry + same-pixel lighting 区段；主指标"),
        ("GPU full-frame time", "µs / frame", "包括 contour/composite；端到端结果"),
        ("CPU command recording", "µs / frame", "只包围 vkBegin→record→vkEnd，不含 present wait"),
        ("Frame-time stability", "median, p95, p99, SD, CV, 1% low", "逐帧 wall time；同时检查 hitch outliers"),
        ("Estimated transactions", "bytes / frame model", "根据 attachment format、extent、load/store/read 次数计算；明确标记 estimated"),
        ("Hardware counters", "tool/vendor-specific", "L2/external memory read/write、tiler bandwidth 等；只作平台内支持证据"),
        ("Implementation cost", "LOC, files, build time, defect count", "辅助讨论 API complexity；不能替代性能结论"),
    ]
    add_table(
        doc,
        ["指标", "单位/统计", "说明"],
        metric_rows,
        [2520, 2700, 4140],
        font_size=8.7,
    )
    add_heading(doc, "6.3 运行协议", 2)
    add_list(
        doc,
        [
            "使用 Release/Benchmark preset；关闭 VSync、动态 UI、asset hot reload、validation 和频繁日志；保留 debug labels 与 timestamps。",
            "每次 run 前记录设备温度/电源状态。Desktop 接电、固定 power plan；Android 固定刷新率、亮度和后台进程，达到温度阈值再开始。",
            "对 path 顺序进行平衡随机化，避免总是让最后一个路径承受热节流。",
            "先运行 300 warm-up frames，再采集 1000 frames；若 shader/pipeline cache 尚未稳定，延长 warm-up 并更新 protocol。",
            "CPU 和 GPU 数据写入独立列；GPU timestamp 需检查 timestampPeriod 与 queue timestampValidBits。",
            "采集 profiler counter 时使用独立 profiling runs；不要把 profiler 注入后的时间与普通 timing runs 混合。",
            "run 结束后写 completion marker 和 file checksum；崩溃/中断数据保留但标为 invalid，不覆盖。",
        ],
        numbered=True,
    )
    add_heading(doc, "6.4 统计与解释", 2)
    add_list(
        doc,
        [
            "先画分布和 time series，检查 warm-up、drift、thermal throttling 与 periodic hitch。",
            "主要报告每个 cell 的 median、p95、p99、SD/CV 和 95% bootstrap confidence interval。",
            "路径比较报告绝对差值与百分比效应；必要时使用配对/重复测量检验，但不只报告 p-value。",
            "若差异小于噪声或 CI 跨过 0，结论写作“未观察到可靠差异”，而不是“完全相同”。",
            "Hardware counters 只支持机制解释；跨 NVIDIA 与 Qualcomm 的 counter 名称/尺度不得直接放在同一数值轴比较。",
            "在论文中分别讨论 correctness、performance、memory-model estimate 与 implementation complexity。",
        ],
    )
    add_heading(doc, "6.5 数据文件结构", 2)
    data_rows = [
        ("raw/<run_id>/frames.csv", "frame index、CPU record、GPU regions、wall frame、present status"),
        ("raw/<run_id>/metadata.json", "commit、config、device/driver/API/features、asset/shader hashes、thermal/power notes"),
        ("raw/<run_id>/capabilities.json", "完整 capability probe snapshot"),
        ("captures/<run_id>/", "RenderDoc/Nsight/AGI capture 或其索引与说明"),
        ("derived/summary.csv", "分析脚本生成的 cell-level summary；可删除重建"),
        ("figures/", "论文图；文件名包含 metric、platform、scene、resolution"),
        ("protocol/protocol-v1.0.md", "冻结后的运行、异常、排除和温控规则"),
    ]
    add_table(
        doc,
        ["路径", "内容"],
        data_rows,
        [3300, 6060],
        font_size=8.8,
    )

    add_heading(doc, "7. 测试与质量保证", 1)
    add_heading(doc, "7.1 测试层级", 2)
    test_rows = [
        ("Unit", "format byte size、alignment、camera interpolation、CSV/JSON serialization、estimated transaction model", "每次 commit / CTest"),
        ("Vulkan smoke", "instance/device/swapchain、resize、device lost handling、Android lifecycle", "每个 build preset"),
        ("Validation", "三路径 × 两 scenes × 两 resolutions，0 error", "每个节点退出前"),
        ("Golden image", "normal/albedo/depth/toon/final screenshots；SSIM/像素阈值", "workload/path 变化时"),
        ("Cross-path equivalence", "相同 frame 的 final image 和关键 attachments 对比", "N3、N5、正式数据前"),
        ("Performance regression", "固定机器短跑；仅检测大幅回退，不替代正式实验", "N7 runner 就绪后每次相关变更"),
        ("Long-run", "10k frames、resize/background/foreground、内存稳定", "N7 前"),
        ("Build reproducibility", "clean clone/build、shader regeneration、asset fetch/hash", "N7、N8"),
    ]
    add_table(
        doc,
        ["层级", "覆盖", "触发"],
        test_rows,
        [1740, 5580, 2040],
        font_size=8.6,
    )
    add_heading(doc, "7.2 三路径公平性审计", 2)
    add_list(
        doc,
        [
            "同一 SPIR-V shader module 或由同一 GLSL source、同一 compiler flags 生成的等价 module。",
            "同一 draw order、mesh、descriptor data、camera matrices、viewport/scissor、sample count。",
            "同一 attachment formats、extent、clear values 与最终可见输出。",
            "允许变化的只有 render organization、attachment access route、必要的 layout/dependency/descriptor binding。",
            "所有 path-specific 差异写入 path_manifest.json；无法解释的差异在正式实验前必须消除。",
        ],
        marker="□",
    )
    add_heading(doc, "7.3 Validation waiver 规则", 2)
    add_body(
        doc,
        "默认不存在“可以忽略的 validation warning”。若确认属于工具/driver 已知问题，必须记录完整 message ID、触发条件、规范依据、"
        "设备/driver、为什么不影响结果和复现步骤，并由 supervisor review 后加入 docs/validation-waivers.md。",
    )

    add_heading(doc, "8. Desktop 与 Android 实施方案", 1)
    add_heading(doc, "8.1 平台共享与隔离", 2)
    platform_rows = [
        ("Window/surface", "GLFW", "GameActivity/ANativeWindow", "仅平台层不同"),
        ("Build", "CMake Preset + Ninja/MSVC", "Gradle + CMake + NDK", "共享 C++ target 与 shaders"),
        ("Assets", "filesystem", "AAssetManager / staged files", "通过 AssetProvider 接口"),
        ("Logging", "console/file", "logcat + file export", "benchmark 模式限制日志"),
        ("Profiling", "Nsight/RenderDoc", "AGI/Qualcomm tool if available", "counter 仅平台内解释"),
        ("Lifecycle", "resize/close", "pause/resume/window recreate", "Android 需独立 stress test"),
    ]
    add_table(
        doc,
        ["关注点", "Desktop", "Android", "统一策略"],
        platform_rows,
        [1560, 2280, 2520, 3000],
        font_size=8.4,
    )
    add_heading(doc, "8.2 DRLR feature probe 决策树", 2)
    add_list(
        doc,
        [
            ("支持 Vulkan 1.4：", "查询 core dynamicRenderingLocalRead，并进一步检查 depth/stencil 与 multisampled local-read properties。"),
            ("仅 Vulkan 1.3：", "枚举并启用 VK_KHR_dynamic_rendering_local_read，查询 VkPhysicalDeviceDynamicRenderingLocalReadFeaturesKHR。"),
            ("Feature 为 false：", "DRLR path 不创建；UI/metadata 显示 Unsupported；该平台实验单元为 NA。"),
            ("Color 支持但 depth local read 受限：", "核心 local-read 段只读取 color normal/albedo；depth 留给公共 contour pass，保持可比。"),
            ("Xiaomi 14 不支持：", "优先借用/更换一台明确支持的 Android 设备；若不可得，保留 desktop 三路径与 mobile multi-pass/subpass 比较，并把限制写入 thesis。"),
        ],
        numbered=True,
    )

    add_heading(doc, "9. 风险登记与应对", 1)
    risk_rows = [
        ("R1", "Xiaomi 14 不支持 DRLR", "高", "高", "在 N5/N6 前完成 feature probe；不等到移植后发现。", "借用兼容设备；否则 mobile cell=NA，重写比较范围。"),
        ("R2", "Sobel 与 local read 语义不兼容", "高", "高", "核心段只做 same-pixel toon；contour 独立公共 pass。", "增加 derivative contour 仅作附加实验，不替换主协议。"),
        ("R3", "直接移植 Afterglow 超时", "高", "中", "采用 clean core；禁止引入 ECS/material/hot reload。", "只复制小型、可审计 MIT 模块。"),
        ("R4", "Android port 晚期失败", "高", "中", "N6 先做 Android triangle，并在路径移植中持续 smoke。", "冻结 desktop 主结果，移动结果作为受限研究。"),
        ("R5", "三路径视觉不等价", "高", "中", "共享 shaders/config；golden attachment diff。", "回退到最小 scene，逐 pass 定位。"),
        ("R6", "GPU timing 噪声/热节流", "高", "中", "warm-up、随机顺序、重复、温控、关闭 VSync。", "延长 cool-down；减少矩阵但保持 repeats。"),
        ("R7", "Counters 无法跨平台比较", "中", "高", "timings 为主；counters 仅平台内支持。", "删除跨平台 counter 数值结论。"),
        ("R8", "Validation/driver bug", "中", "中", "锁定 SDK/driver；保存 message/capture；最小复现。", "更换驱动或记录 waiver/limitation。"),
        ("R9", "资产许可证不清", "中", "低", "优先 procedural/public-domain；保存来源与 hash。", "替换为自制 mesh。"),
        ("R10", "论文写作堆积", "高", "中", "每个节点更新 Methods/Implementation/figures。", "N8 只整合，不从零写。"),
        ("R11", "数据误删/不可复现", "高", "低", "raw append-only、checksum、双份备份、release tags。", "从 runner logs 重跑对应 cells。"),
        ("R12", "范围膨胀", "高", "中", "Must/Should/Could backlog；Could 项不得阻塞当前节点。", "删除 UI 美化、额外效果和额外设备。"),
    ]
    add_table(
        doc,
        ["ID", "风险", "影响", "概率", "预防", "应急"],
        risk_rows,
        [480, 1900, 650, 650, 2740, 2940],
        font_size=7.85,
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
    )
    add_callout(
        doc,
        "停止规则",
        "如果 capability probe 未确认 DRLR support，不能继续假设 Xiaomi 14 可运行 DRLR；如果 N1 的 Multi-pass 视觉基线尚未完成，"
        "不能开始 Subpass 或 DRLR；如果 N7 pilot 数据仍不稳定，不能开始正式实验。",
    )

    add_heading(doc, "10. 开发管理与证据管理", 1)
    add_heading(doc, "10.1 Git 与变更策略", 2)
    add_list(
        doc,
        [
            "主仓库在 workspace root 或 Project 下初始化；AfterglowRender 作为 reference submodule/独立目录，不混入主 history。",
            "分支建议：main（可演示）、develop（整合）、feature/<node-or-slice>；每个节点至少保留一个可运行 tag 或 milestone commit。",
            "commit message 使用类型前缀：core、render、bench、android、test、docs、data。",
            "N3 workload freeze 后，任何 shader/format/scene 更改必须提升 benchmark protocol version 并使旧结果失效。",
            "原始数据不进入普通 Git history；使用 release archive、Git LFS 或独立 data bundle，并保存 checksum manifest。",
        ],
    )
    add_heading(doc, "10.2 Vibe 开发的单节点循环", 2)
    cadence_rows = [
        ("选择", "从当前节点只选一个 vertical slice；写清可见结果和退出条件。"),
        ("生成", "让 AI 生成最小改动；禁止一次同时改 renderer、platform、shader 和数据格式。"),
        ("验证", "立即 build、运行 smoke/validation，并检查画面或导出数据。"),
        ("修正", "根据真实错误和 capture 调整；不要用更多抽象掩盖未理解的问题。"),
        ("留证", "保存 commit、截图/短视频、RenderDoc capture、CSV 或日志路径。"),
        ("整合", "slice 通过后合并；节点通过后创建 tag、备份 docs/data 并冻结输入。"),
    ]
    add_table(
        doc,
        ["步骤", "动作"],
        cadence_rows,
        [2100, 7260],
        font_size=9,
    )
    add_heading(doc, "10.3 Issue 模板", 2)
    add_table(
        doc,
        ["字段", "必须包含"],
        [
            ("Goal", "用户/研究价值；与 objective 或当前节点的关系"),
            ("Scope", "包含与不包含；涉及路径/平台"),
            ("Acceptance", "可运行命令、预期截图/数据/validation 结果"),
            ("Evidence", "commit、capture、CSV、截图或日志路径"),
            ("Risks", "可能影响 fairness、protocol 或 frozen workload 的点"),
        ],
        [1740, 7620],
        font_size=9,
    )

    add_heading(doc, "11. 作品集、论文与演示对齐", 1)
    add_heading(doc, "11.1 Portfolio case study 结构", 2)
    portfolio_rows = [
        ("Hero", "最终场景大图或短视频；一句话说明 Vulkan Multi-pass NPR renderer。"),
        ("Problem", "为什么需要自定义 stylized renderer，以及要解决的 toon、contour 与可调试性问题。"),
        ("My contribution", "明确列出 Vulkan core、G-buffer、toon shader、contour、同步、debug views 和工具链中由你完成的部分。"),
        ("Architecture", "使用简洁架构图说明 Geometry → Toon → Contour → Composite，标出 resources 与 barriers。"),
        ("Visual breakdown", "并列展示 Albedo、Normal、Depth、Toon、Contour 与 Final，不只放最终截图。"),
        ("Technical proof", "RenderDoc capture、validation-clean 状态、GPU timings 和典型 bug/fix；避免没有证据的性能宣传。"),
        ("Reflection", "说明 Afterglow 只作为参考、第三方角色仅作非官方技术研究、哪些选择将留给 FYP benchmark。"),
        ("Access", "公开 Repo/Release、构建说明、视频和资产权利说明；Release 不打包无再分发许可的角色模型与贴图。"),
    ]
    add_table(
        doc,
        ["页面模块", "建议内容"],
        portfolio_rows,
        [2160, 7200],
        font_size=8.9,
    )
    add_heading(doc, "11.2 Thesis evidence map", 2)
    thesis_rows = [
        ("Introduction / Gap", "proposal、feature support truth table、为什么比较三路径"),
        ("Related Work", "Vulkan rendering models、subpass/DRLR、NPR、profiling limitations"),
        ("Methodology", "test matrix、controls、warm-up/repeats、metrics、statistics、ethics/licenses"),
        ("System Design", "Figure 1、interfaces、G-buffer、pass diagrams、platform layer"),
        ("Implementation", "barriers/layouts、subpass dependencies、DRLR mapping、Android differences"),
        ("Results", "timing distributions、effect sizes、estimated transactions、platform counters"),
        ("Discussion", "平台差异、implementation cost、small/no difference 的意义、threats to validity"),
        ("Conclusion", "逐项回答 H1–H3；不超出两台设备与该 workload 的范围"),
    ]
    add_table(
        doc,
        ["章节", "应保存的工程证据"],
        thesis_rows,
        [2460, 6900],
        font_size=8.8,
    )
    add_heading(doc, "11.3 最终演示脚本（建议 3–5 分钟）", 2)
    add_list(
        doc,
        [
            "显示 capability panel：GPU、driver、API、DRLR supported/unsupported。",
            "固定场景运行并切换 debug views：normal、depth、toon、contour。",
            "切换 Multi-pass / Subpass / DRLR，证明视觉结果一致。",
            "进入 benchmark mode，展示计时区段与导出 run ID。",
            "打开一张预生成结果图，解释核心段与 full-frame 结果为何分开。",
            "最后说明平台限制、small/no difference 也属于有效结论。",
        ],
        numbered=True,
    )
    add_heading(doc, "11.4 最终提交包", 2)
    add_list(
        doc,
        [
            "Source release + pinned dependencies + build instructions。",
            "Windows executable 与 Android APK（如学校允许提交 binary）。",
            "Protocol、capabilities、sample raw data、analysis scripts、derived figures。",
            "Documentation、licenses、asset attributions；thesis、slides/poster、demo video 与 FAQ。",
        ],
        marker="□",
    )

    add_heading(doc, "12. 立即执行顺序：作品集优先", 1)
    add_callout(
        doc,
        "第一目标",
        "先完成一条高完成度、可公开展示的 Desktop Multi-pass NPR renderer。"
        "在这个作品集节点通过前，不同时开发 Subpass、DRLR 或 Android。",
    )
    startup_rows = [
        ("S0", "记录已确认决定：Multi-pass、类《终末地》角色/工业场景、公开代码、最小 ImGui、Afterglow reference-only；登记 UE 5.7 莱万汀资产位置与授权边界。", "U01–U08 decision log 冻结"),
        ("S1", "整理目录并初始化主 repo；建立 Portfolio/Debug presets、README、LICENSES、assets_private/.gitignore 与公开 placeholder。", "clean build 可重复且私有资产不会被提交"),
        ("S2", "从 MyVulkanApp 完成 instance、validation、device/queue、GLFW surface、swapchain、frames-in-flight 与 resize。", "稳定清屏/triangle"),
        ("S3", "实现最小 RAII 资源层和 shader build；加入 object names、debug labels 与 RenderDoc capture 入口。", "资源生命周期 validation clean"),
        ("S4", "在 UE 5.7 中导出莱万汀 mesh/纹理与 manifest；首版使用静态姿态，完成 material ID/mask 映射、alpha cutout、two-sided 与基础灯光。", "导出清单完整；角色和场景材质槽正确"),
        ("S5", "实现 G-buffer Geometry pass，并加入 normal/albedo/depth debug views。", "attachments 经 RenderDoc 验证"),
        ("S6", "实现材质感知的 Multi-pass toon：skin/hair/fabric/metal ramps、高光、rim light 与 emissive mask。", "角色半身与全身画面稳定"),
        ("S7", "实现轻量 normal/depth contour 与 final composite；调试深度、发丝边缘、轮廓强度、tone mapping 和色彩空间。", "目标视觉达到作品集标准"),
        ("S8", "加入最小 ImGui：相机、灯光、材质参数、toon/contour、debug views、GPU overlay；完成稳定性测试。", "Portfolio preset 功能冻结"),
        ("S9", "制作 Release、README、架构图、分解截图、短视频、技术说明与许可证包；确认公开包不含受限模型/贴图。", "N2 Portfolio ready"),
        ("S10", "从作品集版本提取共享 workload 与 RenderPath 接口，加入 capability_probe、timestamps 和 CSV/JSON；再进入研究路径。", "N3 Benchmark-ready"),
    ]
    add_table(
        doc,
        ["顺序", "任务", "退出条件"],
        startup_rows,
        [1020, 5940, 2400],
        font_size=8.6,
    )
    add_heading(doc, "12.1 能力探测后的 Go / No-Go 决策", 2)
    add_table(
        doc,
        ["情况", "决定"],
        [
            ("两平台都支持 DRLR", "执行原始 3×2 平台矩阵；优先 color local-read，depth 属性单独记录。"),
            ("仅 desktop 支持 DRLR", "Desktop 做三路径主实验；Mobile 做 multi-pass/subpass 平台补充；修改 H2 表述。"),
            ("两平台都不支持 DRLR", "优先获得兼容设备；否则将 FYP 改为 implementation/feasibility + subpass benchmark，并立即与 supervisor 批准 scope change。"),
            ("Android 基础移植失败", "N6 先修复 platform shell；禁止在移动端先做 GUI、复杂资产系统或额外 NPR 效果。"),
        ],
        [3000, 6360],
        font_size=8.8,
    )

    add_heading(doc, "附录 A · AfterglowRender 借鉴清单", 1)
    afterglow_rows = [
        ("AfterglowRenderPass / AfterglowSubpassContext", "subpass attachments、references、dependencies 的组织方式", "重写为小型 C++17 class；不复制 proxy/reflection 基类"),
        ("AfterglowFramebufferManager", "attachment usage 与 framebuffer 对应关系", "只借鉴 validation 思路；三路径统一由 GBufferLayout 创建"),
        ("AfterglowMaterialAsset / Manager", "input attachment 声明与 descriptor 绑定概念", "不复用 runtime material system；改为固定 descriptor layouts"),
        ("AfterglowCommandBuffer / Synchronizer", "command 生命周期与同步命名", "用直接 Vulkan + Sync2 helpers 重写"),
        ("AfterglowDevice / PhysicalDevice", "extension screening", "必须升级为 Features2 pNext，不能沿用只查询 VkPhysicalDeviceFeatures 的方式"),
        ("Afterglow Shaders/Toon", "视觉参考", "不直接作为研究 shader；重写统一 GLSL"),
        ("Afterglow GUI / hot reload / ECS", "无", "排除在 FYP core 外"),
    ]
    add_table(
        doc,
        ["参考位置", "可借鉴", "边界"],
        afterglow_rows,
        [3300, 2760, 3300],
        font_size=8.5,
    )
    add_body(
        doc,
        "Reference snapshot: AfterglowRender commit 8380690 (2026-05-03), origin "
        "https://github.com/Celumina/AfterglowRender, MIT License.",
        italic=True,
        color=MUTED,
    )

    add_heading(doc, "附录 B · Benchmark 配置建议", 1)
    config_rows = [
        ("path", "multipass | subpass | drlr"),
        ("scene", "low | high"),
        ("resolution", "1280x720 | 1920x1080"),
        ("warmup_frames", "300"),
        ("sample_frames", "1000"),
        ("repeat_index", "0–9"),
        ("frames_in_flight", "2（所有路径相同）"),
        ("vsync", "false"),
        ("ui", "false in Benchmark preset"),
        ("validation", "false for timing; true for validation run"),
        ("camera_script", "assets/camera/fixed_path_v1.json"),
        ("protocol_version", "1.0 after pilot freeze"),
        ("seed", "固定且写入 metadata"),
    ]
    add_table(
        doc,
        ["Key", "Value / Rule"],
        config_rows,
        [3000, 6360],
        font_size=8.9,
    )

    add_heading(doc, "附录 C · Node exit checklist", 1)
    checklist_groups = [
        (
            "N0 · Foundation",
            [
                "RTX 4060 capability JSON saved",
                "Xiaomi 14 capability JSON saved",
                "DRLR core/extension/feature verdict recorded",
                "Depth/stencil and MSAA local-read properties recorded",
                "Profiler availability recorded",
                "Plan A/B approved",
            ],
        ),
        (
            "N2 · Portfolio ready",
            [
                "Windows Release build verified",
                "Portfolio scene and final screenshots frozen",
                "G-buffer / Toon / Contour debug views available",
                "README and architecture figure complete",
                "Demo video and technical summary complete",
                "Public assets and licenses audited",
            ],
        ),
        (
            "N3 · Workload freeze",
            [
                "Two scenes and two resolutions render",
                "Camera deterministic",
                "G-buffer formats frozen",
                "Toon and contour parameters frozen",
                "Golden images and shader hashes saved",
                "Contour separated from local-read comparison",
            ],
        ),
        (
            "N7 · Pilot ready",
            [
                "All supported paths validation-clean",
                "Cross-path image diff passed",
                "Runner survives interrupted run",
                "CSV/JSON schema validated",
                "Timestamp calibration verified",
                "Thermal/power SOP rehearsed",
            ],
        ),
        (
            "N8 · Final",
            [
                "Release builds verified on clean machines/devices",
                "Raw data checksum manifest present",
                "Figures reproducible from scripts",
                "Licenses and citations complete",
                "Thesis claims map to evidence",
                "Offline demo video available",
            ],
        ),
    ]
    for heading, items in checklist_groups:
        add_heading(doc, heading, 2)
        add_list(doc, items, marker="□", compact=True)

    add_heading(doc, "附录 D · 主要技术来源", 1)
    sources = [
        (
            "Project proposal",
            "DMT2309242-WuChenfeng-Proposal.docx（workspace source of truth for research scope）",
            None,
        ),
        (
            "AfterglowRender",
            "Local source snapshot and MIT license",
            "https://github.com/Celumina/AfterglowRender",
        ),
        (
            "VK_KHR_dynamic_rendering_local_read",
            "Extension reference and Vulkan 1.4 promotion notes",
            "https://docs.vulkan.org/refpages/latest/refpages/source/VK_KHR_dynamic_rendering_local_read.html",
        ),
        (
            "Khronos DRLR sample",
            "Side-by-side dynamic rendering local read and subpass implementation",
            "https://docs.vulkan.org/samples/latest/samples/extensions/dynamic_rendering_local_read/README.html",
        ),
        (
            "Vulkan Render Pass specification",
            "Input attachments read the same framebuffer location",
            "https://docs.vulkan.org/spec/latest/chapters/renderpass.html",
        ),
        (
            "Android GPU Inspector",
            "Android system and frame profiling",
            "https://developer.android.com/agi",
        ),
        (
            "GRYPHLINE Account Terms of Service",
            "Public repository boundary for game content and intellectual-property restrictions",
            "https://user.gryphline.com/en-us/protocol/terms_of_service",
        ),
    ]
    for title, note, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{title}. ")
        set_run(r, size=10, bold=True, color=NAVY)
        r = p.add_run(note)
        set_run(r, size=10, color=INK)
        if url:
            p.add_run(" ")
            add_hyperlink(p, "Open source", url)

    add_callout(
        doc,
        "下一步",
        "U01–U08 已确认，莱万汀作为首个私有测试角色。下一步在 UE 5.7 中核对 Skeletal Mesh 的材质槽、LOD、"
        "贴图通道与色彩空间，生成一次可重复的私有导出包和 manifest；随后执行 S1–S9 完成 Desktop Multi-pass "
        "作品集 renderer，N2 通过后再进入研究路径。",
    )

    # Core properties and metadata
    props = doc.core_properties
    props.title = "FYP Development Plan"
    props.subject = "Vulkan Subpasses vs Dynamic Rendering Local Read for Real-Time NPR"
    props.author = "Wu Chenfeng"
    props.keywords = "FYP, Vulkan, DRLR, Subpasses, NPR, Benchmark"
    props.comments = (
        "Generated from the proposal and local AfterglowRender review. "
        "Chinese East Asian font override: Microsoft YaHei."
    )

    # Keep all sections consistent.
    for section in doc.sections:
        section.page_width = Twips(PAGE_WIDTH_DXA)
        section.page_height = Twips(PAGE_HEIGHT_DXA)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(f"Saved: {output}")
    issues = audit_docx_tables(output)
    if issues:
        raise SystemExit(f"Table geometry audit failed with {issues} issue(s)")
